"""Unit tests for Google Drive menu parsing (apps/food/services/drive_menu.py).

These build .docx files in memory and call ``DriveMenuService._parse_docx``
directly, so they need no network access, Google Drive credentials, or database.
"""

import io

import pytest
from docx import Document
from docx.oxml import OxmlElement

from apps.food.services.drive_menu import DriveMenuService


def _make_docx(lines: list[str]) -> io.BytesIO:
    """Build an in-memory .docx with one paragraph per given line."""
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _make_docx_with_pagebreak(page1_lines: list[str], page2_lines: list[str]) -> io.BytesIO:
    """Build a .docx whose content is split across a page boundary.

    The first paragraph of ``page2_lines`` carries a ``w:lastRenderedPageBreak``
    marker — exactly what Word/Google Docs write when text flows onto a new page
    automatically, which is the real-world cause of the missing-Thursday bug. A
    manual page break is added too, so the break is detectable regardless of how
    ``_has_page_break`` looks for it.
    """
    doc = Document()
    for line in page1_lines:
        doc.add_paragraph(line)
    doc.add_page_break()
    for i, line in enumerate(page2_lines):
        para = doc.add_paragraph()
        run = para.add_run(line)
        if i == 0:
            run._r.insert(0, OxmlElement("w:lastRenderedPageBreak"))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@pytest.fixture
def service() -> DriveMenuService:
    return DriveMenuService()


def test_parses_simple_overview(service: DriveMenuService) -> None:
    """A plain overview with each day once is parsed for all four days."""
    parsed = service._parse_docx(
        _make_docx(
            [
                "Mandag",
                "Lasagne",
                "Tirsdag",
                "Thai karry",
                "Onsdag",
                "Frikadeller",
                "Torsdag",
                "Pasta pesto",
            ]
        ),
        week_number=10,
    )
    assert parsed.monday == "Lasagne"
    assert parsed.tuesday == "Thai karry"
    assert parsed.wednesday == "Frikadeller"
    assert parsed.thursday == "Pasta pesto"


def test_thursday_dish_spilled_to_next_page(service: DriveMenuService) -> None:
    """Regression for the reported bug: Torsdag's header sits at the very bottom
    of page 1 and its dish flows onto page 2.

    The parser reads only page 1 (recipes live on later pages), so the dish was
    dropped and Thursday showed up empty. The fix rescues the first paragraph of
    the next page as the dangling day's dish — while still excluding the recipe
    text that follows it.
    """
    parsed = service._parse_docx(
        _make_docx_with_pagebreak(
            page1_lines=[
                "Mandag",
                "Majsdeller med nye kartofler",
                "Tirsdag",
                "Grøn karry med ris",
                "Onsdag",
                "Köfte med tabouleh",
                "Torsdag",  # header is the last line on page 1
            ],
            page2_lines=[
                "Hvid kartoffelpizza med løgkompot",  # the dish, top of page 2
                "Lang opskrift som ikke må ende i torsdagsretten",  # recipe text
                "Mandag",
                "Detaljeret opskrift for mandag",
            ],
        ),
        week_number=23,
    )
    assert parsed.monday == "Majsdeller med nye kartofler"
    assert parsed.tuesday == "Grøn karry med ris"
    assert parsed.wednesday == "Köfte med tabouleh"
    # The dish is rescued from the top of page 2...
    assert parsed.thursday == "Hvid kartoffelpizza med løgkompot"
    # ...but the recipe text after it must NOT leak in (proves the page break is
    # still respected — otherwise the whole of page 2 would be appended).
    assert "Lang opskrift" not in parsed.thursday


def test_dish_on_page_one_is_not_overwritten_by_page_two(service: DriveMenuService) -> None:
    """When Thursday's dish already fits on page 1, the next page's content (the
    recipe section) must not replace it."""
    parsed = service._parse_docx(
        _make_docx_with_pagebreak(
            page1_lines=[
                "Mandag",
                "Marokkansk gryderet",
                "Tirsdag",
                "Aubergine ragout",
                "Onsdag",
                "Boller i karry",
                "Torsdag",
                "Sprøde gulerods nuggets",  # dish present on page 1
            ],
            page2_lines=[
                "Mandag",
                "Detaljeret opskrift for mandag",
            ],
        ),
        week_number=25,
    )
    assert parsed.monday == "Marokkansk gryderet"
    assert parsed.tuesday == "Aubergine ragout"
    assert parsed.wednesday == "Boller i karry"
    assert parsed.thursday == "Sprøde gulerods nuggets"


def test_weekly_veggie_order_not_appended_to_thursday(service: DriveMenuService) -> None:
    """Regression for the reported bug: the weekly grocery/veggie order
    ("Gnavegrønt til ugen ...") sits after Torsdag's dish with no weekday header,
    so it used to be swallowed into Thursday's menu text. It must now be excluded
    while Thursday's real dish is kept and Mon-Wed are unaffected.
    """
    parsed = service._parse_docx(
        _make_docx(
            [
                "Mandag",
                "Lasagne",
                "Tirsdag",
                "Thai karry",
                "Onsdag",
                "Frikadeller",
                "Torsdag",
                "Beluga-bolognese med spaghetti og fennikelsalat med citrusfrugter",
                "Gnavegrønt til ugen 18 stk agurker 4 stk blomkål "
                "6,8 kg gulerødder 5 kg appelsiner 10 stk glaskål",
            ]
        ),
        week_number=24,
    )
    assert parsed.monday == "Lasagne"
    assert parsed.tuesday == "Thai karry"
    assert parsed.wednesday == "Frikadeller"
    assert parsed.thursday == "Beluga-bolognese med spaghetti og fennikelsalat med citrusfrugter"
    assert "Gnavegrønt" not in parsed.thursday
    assert "agurker" not in parsed.thursday


@pytest.mark.parametrize(
    "veggie_order_header",
    [
        "Gnavegrønt til ugen:",
        "Gnavegrønt til ugen",
        "Gnavegrønt.",
        "Gnavegrønt\ttil ugen:",  # tab instead of space (uge 20)
        "Gnavegrønt: Gulerødder, æbler, pære, glaskål og blomkål",
        "Grønt til ugen",
        "Ugens grønt",
    ],
)
def test_veggie_order_header_variants_all_terminate(
    service: DriveMenuService, veggie_order_header: str
) -> None:
    """Every phrasing of the shopping-block header seen in the menu archive must
    end Thursday's dish text."""
    parsed = service._parse_docx(
        _make_docx(
            [
                "Torsdag",
                "Cremet blomkålssuppe med foccacia",
                veggie_order_header,
                "æbler 4 kg",
                "gulerødder 7 kg",
            ]
        ),
        week_number=8,
    )
    assert parsed.thursday == "Cremet blomkålssuppe med foccacia"
    assert "kg" not in parsed.thursday


def test_gnavegroent_as_a_days_side_dish_does_not_stop_parsing(
    service: DriveMenuService,
) -> None:
    """Cooks list gnavegrønt as a side dish for individual days ("Tilbehør: ... +
    gnavegrønt"). Those lines belong to their day and must not terminate parsing.

    Taken verbatim from uge 5/2026. Treating them as the weekly shopping block
    stopped the parse at Monday and blanked Tuesday through Thursday.
    """
    parsed = service._parse_docx(
        _make_docx(
            [
                "Uge med Dhal",
                "Mandag",
                "Svampe i butterdej med porre og kål",
                "Tilbehør: bulgursalat med kål og kerner + gnavegrønt",
                "Tirsdag",
                "Dhal med raita, bagte rodfrugter",
                "Tilbehør: Flutes og gnavegrønt",
                "Onsdag",
                "Pitabrød med BBQ-svin, grønt, urtedressing og kartofler",
                "Tilbehør: brød og gnavegrønt",
                "Torsdag",
                "Græskinspireret 'byg-selv' pastasalat",
                "Tilbehør: brød og gnavegrønt",
            ]
        ),
        week_number=5,
    )
    assert parsed.monday == (
        "Svampe i butterdej med porre og kål Tilbehør: bulgursalat med kål og kerner + gnavegrønt"
    )
    assert parsed.tuesday == "Dhal med raita, bagte rodfrugter Tilbehør: Flutes og gnavegrønt"
    assert parsed.wednesday.startswith("Pitabrød med BBQ-svin")
    assert parsed.thursday.startswith("Græskinspireret")
    assert all(day for day in (parsed.tuesday, parsed.wednesday, parsed.thursday))


def test_veggie_order_before_a_later_day_keeps_that_day(service: DriveMenuService) -> None:
    """A weekday header after the shopping block is still parsed.

    The terminator skips the block's lines instead of abandoning the document, so
    a misplaced block can never cost more than the day it interrupts.
    """
    parsed = service._parse_docx(
        _make_docx(
            [
                "Mandag",
                "Lasagne",
                "Gnavegrønt til ugen",
                "agurker 6 kg",
                "Torsdag",
                "Pasta pesto",
            ]
        ),
        week_number=12,
    )
    assert parsed.monday == "Lasagne"
    assert parsed.thursday == "Pasta pesto"
    assert "agurker" not in parsed.monday


def test_repeated_earlier_day_does_not_drop_later_days(service: DriveMenuService) -> None:
    """If an earlier day header repeats (e.g. overview + detailed section on the
    same page) before later days appear, the later days must still be parsed.

    The parser used to ``break`` on the first repeated header, dropping every
    not-yet-seen day. Now it keeps the first occurrence and keeps going.
    """
    parsed = service._parse_docx(
        _make_docx(
            [
                "Mandag",
                "Lasagne",
                "Tirsdag",
                "Thai karry",
                # Repeated earlier headers BEFORE Onsdag/Torsdag first appear
                "Mandag",
                "Opskrift: lasagne",
                "Tirsdag",
                "Opskrift: karry",
                "Onsdag",
                "Frikadeller",
                "Torsdag",
                "Pasta pesto",
            ]
        ),
        week_number=11,
    )
    assert parsed.monday == "Lasagne"
    assert parsed.tuesday == "Thai karry"
    assert parsed.wednesday == "Frikadeller", "Wednesday must not be dropped"
    assert parsed.thursday == "Pasta pesto", "Thursday must not be dropped"


def test_page2_rescue_keeps_a_dish_that_lists_gnavegroent_as_a_side(
    service: DriveMenuService,
) -> None:
    """Torsdag's dish must survive spilling onto page 2 with a gnavegrønt side.

    The rescue used to test the overflow paragraph with ``.search``, so any dish
    mentioning gnavegrønt anywhere was read as the weekly shopping block and
    discarded — leaving Torsdag blank. The main loop anchors the same pattern
    with ``.match`` precisely because those mid-line mentions are ordinary menu
    text and outnumber the real shopping blocks in the archive.
    """
    parsed = service._parse_docx(
        _make_docx_with_pagebreak(
            page1_lines=[
                "Mandag",
                "Lasagne",
                "Tirsdag",
                "Frikadeller",
                "Onsdag",
                "Fisk",
                "Torsdag",
            ],
            page2_lines=[
                "Butter chicken med ris. Tilbehør: bulgursalat med kål og kerner + gnavegrønt"
            ],
        ),
        week_number=4,
    )
    assert parsed.thursday.startswith("Butter chicken"), "Thursday's dish was discarded"


def test_page2_rescue_still_drops_the_weekly_shopping_block(
    service: DriveMenuService,
) -> None:
    """A page-2 paragraph that *begins* with the shopping block is not a dish."""
    parsed = service._parse_docx(
        _make_docx_with_pagebreak(
            page1_lines=[
                "Mandag",
                "Lasagne",
                "Tirsdag",
                "Frikadeller",
                "Onsdag",
                "Fisk",
                "Torsdag",
            ],
            page2_lines=["Gnavegrønt til ugen: 18 stk agurker, 2 kg gulerødder"],
        ),
        week_number=4,
    )
    assert parsed.thursday == "", "the shopping block must not become Thursday's dish"
