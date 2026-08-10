"""
Tests for the Food app.

Uses pytest and pytest-django for testing.
"""

from datetime import date, datetime, time, timedelta
from decimal import Decimal
from unittest.mock import patch

import pytest
from django.urls import reverse
from django.utils import timezone

from apps.food.models import (
    CycleStatus,
    DayOfWeek,
    DiningOption,
    DriveMenuCache,
    FoodTeam,
    FoodTeamCycle,
    FoodTeamMember,
    FoodTeamWish,
    FoodTicket,
    MealPreference,
    MealRegistration,
    SeatingTime,
    SwapRequestStatus,
    TeamSwapRequest,
)
from apps.food.serializers import (
    FoodTeamCycleCreateSerializer,
    FoodTeamWishCreateUpdateSerializer,
    FoodTicketCreateSerializer,
    MealRegistrationCreateUpdateSerializer,
)
from apps.food.services.team_generator import TeamGenerationResult, TeamGenerator
from apps.users.models import User
from conftest import generate_cooking_dates

# =============================================================================
# Model Tests
# =============================================================================


class TestMealPreferenceModel:
    """Tests for MealPreference model."""

    def test_create_meal_preference(self, meal_preference):
        """Test creating a meal preference."""
        assert meal_preference.day_of_week == DayOfWeek.MONDAY
        assert meal_preference.adults_meat == 0
        assert meal_preference.adults_veg == 2
        assert meal_preference.adults_count == 2
        assert meal_preference.children_count == 1
        assert meal_preference.dining_option == DiningOption.EAT_IN

    def test_meal_preference_unique_constraint(self, db, house):
        """Test unique constraint on house + day_of_week."""
        from django.db import IntegrityError

        MealPreference.objects.create(
            house=house,
            day_of_week=DayOfWeek.MONDAY,
        )
        with pytest.raises(IntegrityError):
            MealPreference.objects.create(
                house=house,
                day_of_week=DayOfWeek.MONDAY,
            )


class TestMealRegistrationModel:
    """Tests for MealRegistration model."""

    def test_create_meal_registration(self, meal_registration):
        """Test creating a meal registration."""
        assert meal_registration.adults_meat == 0
        assert meal_registration.adults_veg == 2
        assert meal_registration.adults_count == 2
        assert meal_registration.children_count == 1
        assert meal_registration.is_active is True

    def test_total_portions(self, meal_registration):
        """Test total_portions property."""
        assert meal_registration.total_portions == 3  # 2 adults + 1 child

    def test_meal_registration_with_house(self, db, user_with_house, monday_date, house):
        """Test meal registration with house."""
        reg = MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=monday_date,
            adults_veg=4,
        )
        assert reg.house == house
        assert "House 1" in str(reg)


class TestFoodTicketModel:
    """Tests for FoodTicket model."""

    def test_create_food_ticket(self, food_ticket):
        """Test creating a food ticket."""
        assert food_ticket.price == Decimal("50.00")
        assert food_ticket.is_available is True
        assert food_ticket.is_free is False

    def test_free_ticket(self, food_ticket_free):
        """Test free food ticket."""
        assert food_ticket_free.is_free is True
        assert food_ticket_free.total_portions == 2

    def test_claim_ticket(self, db, food_ticket, admin_user):
        """Test claiming a food ticket."""
        food_ticket.is_available = False
        food_ticket.claimed_by = admin_user
        food_ticket.claimed_at = timezone.now()
        food_ticket.save()

        assert food_ticket.is_available is False
        assert food_ticket.claimed_by == admin_user


class TestFoodTeamCycleModel:
    """Tests for FoodTeamCycle model."""

    def test_create_cycle(self, food_team_cycle):
        """Test creating a food team cycle."""
        assert food_team_cycle.status == CycleStatus.COLLECTING_WISHES
        assert food_team_cycle.is_accepting_wishes is True

    def test_cooking_dates(self, food_team_cycle):
        """Test cooking_dates contains valid date strings."""
        dates = food_team_cycle.cooking_dates
        assert len(dates) > 0
        for d in dates:
            # Dates are now ISO format strings
            parsed = date.fromisoformat(d)
            assert parsed.weekday() <= 3  # Mon=0, Thu=3

    def test_cycle_not_accepting_after_deadline(self, db, admin_user, monday_date):
        """Test cycle stops accepting wishes after deadline."""
        cooking_dates = generate_cooking_dates(monday_date, num_weeks=4)
        cycle = FoodTeamCycle.objects.create(
            name="Past Deadline Cycle",
            cooking_dates=cooking_dates,
            wish_deadline=timezone.now() - timedelta(days=1),  # Past deadline
            status=CycleStatus.COLLECTING_WISHES,
            created_by=admin_user,
        )
        assert cycle.is_accepting_wishes is False


class TestFoodTeamModel:
    """Tests for FoodTeam model."""

    def test_create_food_team(self, food_team):
        """Test creating a food team."""
        assert food_team.date is not None
        assert food_team.day_name in [
            "Mandag",
            "Tirsdag",
            "Onsdag",
            "Torsdag",
            "Fredag",
            "Lørdag",
            "Søndag",
        ]

    def test_member_count(self, food_team, food_team_member):
        """Test member_count property."""
        assert food_team.member_count == 1


class TestTeamSwapRequest:
    """Tests for TeamSwapRequest model."""

    def test_create_swap_request(self, db, food_team, user, admin_user):
        """Test creating a team swap request."""
        member1 = FoodTeamMember.objects.create(team=food_team, user=user, house_number="1")

        team2 = FoodTeam.objects.create(
            cycle=food_team.cycle,
            date=food_team.date + timedelta(days=1),
        )
        member2 = FoodTeamMember.objects.create(team=team2, user=admin_user, house_number="2")

        swap = TeamSwapRequest.objects.create(
            requester=user,
            requester_membership=member1,
            target_membership=member2,
            message="Please swap with me",
        )

        assert swap.status == SwapRequestStatus.PENDING
        assert swap.target_user == admin_user


class TestDriveMenuCacheModel:
    """Tests for DriveMenuCache model."""

    def test_create_drive_menu_cache(self, db):
        """Test creating a drive menu cache entry."""
        cache = DriveMenuCache.objects.create(
            week_number=3,
            year=2026,
            monday_menu="Lasagne",
            tuesday_menu="Thai curry",
            wednesday_menu="Frikadeller",
            thursday_menu="Pasta",
        )
        assert cache.week_number == 3
        assert cache.year == 2026
        assert str(cache) == "Week 3, 2026"

    def test_is_stale(self, db):
        """Test is_stale method."""
        cache = DriveMenuCache.objects.create(
            week_number=4,
            year=2026,
            monday_menu="Test",
        )
        # Just created, should not be stale
        assert cache.is_stale(max_age_hours=1) is False

    def test_unique_constraint(self, db):
        """Test unique constraint on week_number + year."""
        from django.db import IntegrityError

        DriveMenuCache.objects.create(week_number=5, year=2026)
        with pytest.raises(IntegrityError):
            DriveMenuCache.objects.create(week_number=5, year=2026)


# =============================================================================
# Serializer Tests
# =============================================================================


class MockRequest:
    """Simple mock request object for serializer tests."""

    def __init__(self, user):
        self.user = user


class TestMealRegistrationSerializer:
    """Tests for MealRegistrationCreateUpdateSerializer."""

    def test_validate_weekday(self, db, user_with_house, monday_date):
        """Test that only Mon-Thu dates are accepted."""
        # Valid Monday
        serializer = MealRegistrationCreateUpdateSerializer(
            data={
                "date": monday_date.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
            },
            context={"request": MockRequest(user_with_house)},
        )
        assert serializer.is_valid()

    def test_reject_weekend(self, db, user_with_house, monday_date):
        """Test that weekend dates are rejected."""
        saturday = monday_date + timedelta(days=5)
        serializer = MealRegistrationCreateUpdateSerializer(
            data={
                "date": saturday.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
            },
            context={"request": MockRequest(user_with_house)},
        )
        assert not serializer.is_valid()
        assert "date" in serializer.errors

    def test_reject_meat_on_non_wednesday(self, db, user_with_house, monday_date):
        """Test that adults_meat > 0 on non-Wednesday is rejected."""
        serializer = MealRegistrationCreateUpdateSerializer(
            data={
                "date": monday_date.isoformat(),
                "adults_meat": 1,
                "adults_veg": 0,
                "children_count": 0,
            },
            context={"request": MockRequest(user_with_house)},
        )
        assert not serializer.is_valid()
        assert "adults_meat" in serializer.errors

    def test_reject_empty_portions(self, db, user_with_house, monday_date):
        """Test that zero total portions is rejected when active."""
        serializer = MealRegistrationCreateUpdateSerializer(
            data={
                "date": monday_date.isoformat(),
                "adults_meat": 0,
                "adults_veg": 0,
                "children_count": 0,
                "is_active": True,
            },
            context={"request": MockRequest(user_with_house)},
        )
        assert not serializer.is_valid()

    def test_wednesday_allows_meat(self, db, user_with_house, monday_date):
        """Test that adults_meat > 0 is allowed on Wednesday."""
        wednesday = monday_date + timedelta(days=2)
        serializer = MealRegistrationCreateUpdateSerializer(
            data={
                "date": wednesday.isoformat(),
                "adults_meat": 1,
                "adults_veg": 1,
                "children_count": 0,
            },
            context={"request": MockRequest(user_with_house)},
        )
        assert serializer.is_valid()


class TestFoodTicketCreateSerializer:
    """Tests for FoodTicketCreateSerializer."""

    def test_reject_past_date(self, db, user):
        """Test that past dates are rejected."""
        past_date = timezone.now().date() - timedelta(days=1)
        serializer = FoodTicketCreateSerializer(
            data={
                "date": past_date.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
            },
            context={"request": MockRequest(user)},
        )
        assert not serializer.is_valid()
        assert "date" in serializer.errors


class TestFoodTeamCycleCreateSerializer:
    """Tests for FoodTeamCycleCreateSerializer."""

    def test_valid_cooking_dates(self, db, user, monday_date):
        """Test valid cooking dates are accepted."""
        cooking_dates = [
            monday_date.isoformat(),
            (monday_date + timedelta(days=1)).isoformat(),
            (monday_date + timedelta(days=2)).isoformat(),
        ]
        serializer = FoodTeamCycleCreateSerializer(
            data={
                "name": "Test Cycle",
                "cooking_dates": cooking_dates,
                "wish_deadline": (timezone.now() + timedelta(days=1)).isoformat(),
            },
            context={"request": MockRequest(user)},
        )
        assert serializer.is_valid()

    def test_empty_cooking_dates_rejected(self, db, user):
        """Test empty cooking dates are rejected."""
        serializer = FoodTeamCycleCreateSerializer(
            data={
                "name": "Test Cycle",
                "cooking_dates": [],
                "wish_deadline": (timezone.now() + timedelta(days=1)).isoformat(),
            },
            context={"request": MockRequest(user)},
        )
        assert not serializer.is_valid()
        assert "cooking_dates" in serializer.errors

    def test_cooking_dates_sorted(self, db, user, monday_date):
        """Test cooking dates are sorted in output."""
        # Input dates in random order
        cooking_dates = [
            (monday_date + timedelta(days=2)).isoformat(),
            monday_date.isoformat(),
            (monday_date + timedelta(days=1)).isoformat(),
        ]
        serializer = FoodTeamCycleCreateSerializer(
            data={
                "name": "Test Cycle",
                "cooking_dates": cooking_dates,
                "wish_deadline": (timezone.now() + timedelta(days=1)).isoformat(),
            },
            context={"request": MockRequest(user)},
        )
        assert serializer.is_valid()
        # Dates should be sorted
        result = serializer.validated_data["cooking_dates"]
        assert result == sorted(result)


class TestFoodTeamWishSerializer:
    """Tests for FoodTeamWishCreateUpdateSerializer."""

    def test_validate_dates(self, db, food_team_cycle, user, monday_date):
        """Test date validation."""
        available_dates = [
            monday_date.isoformat(),
            (monday_date + timedelta(days=1)).isoformat(),
        ]
        serializer = FoodTeamWishCreateUpdateSerializer(
            data={
                "cycle": food_team_cycle.id,
                "available_dates": available_dates,
                "comment": "Test",
            },
            context={"request": MockRequest(user)},
        )
        assert serializer.is_valid()

    def test_reject_cycle_not_accepting(self, db, admin_user, user, monday_date):
        """Test rejection when cycle not accepting wishes."""
        cooking_dates = generate_cooking_dates(monday_date, num_weeks=1)
        cycle = FoodTeamCycle.objects.create(
            name="Closed Cycle",
            cooking_dates=cooking_dates,
            wish_deadline=timezone.now() - timedelta(days=1),
            status=CycleStatus.COLLECTING_WISHES,
            created_by=admin_user,
        )
        serializer = FoodTeamWishCreateUpdateSerializer(
            data={
                "cycle": cycle.id,
                "available_dates": [monday_date.isoformat()],
            },
            context={"request": MockRequest(user)},
        )
        assert not serializer.is_valid()


# =============================================================================
# View/API Tests
# =============================================================================


@pytest.mark.django_db
class TestMealPreferenceViews:
    """Tests for meal preference API endpoints."""

    def test_list_preferences(self, authenticated_client, meal_preference):
        """Test listing user's preferences."""
        url = reverse("food:preference-list")
        response = authenticated_client.get(url)
        assert response.status_code == 200

    def test_create_preference(self, api_client, user_with_house):
        """Test creating a meal preference."""
        api_client.force_authenticate(user=user_with_house)
        url = reverse("food:preference-list")
        data = {
            "day_of_week": DayOfWeek.TUESDAY,
            "adults_meat": 0,
            "adults_veg": 2,
            "children_count": 0,
            "dining_option": DiningOption.TAKE_AWAY,
            "seating_time": SeatingTime.SECOND,
        }
        response = api_client.post(url, data, format="json")
        assert response.status_code == 201


@pytest.mark.django_db
class TestMealRegistrationViews:
    """Tests for meal registration API endpoints."""

    def test_list_registrations(self, authenticated_client, meal_registration):
        """Test listing user's registrations."""
        url = reverse("food:registration-list")
        response = authenticated_client.get(url)
        assert response.status_code == 200

    def test_create_registration(self, api_client, user_with_house, monday_date):
        """Test creating a meal registration."""
        api_client.force_authenticate(user=user_with_house)
        url = reverse("food:registration-list")
        # Use next week to avoid duplicates
        next_monday = monday_date + timedelta(weeks=2)
        data = {
            "date": next_monday.isoformat(),
            "adults_meat": 0,
            "adults_veg": 2,
            "children_count": 1,
            "dining_option": DiningOption.EAT_IN,
            "seating_time": SeatingTime.FIRST,
        }
        response = api_client.post(url, data, format="json")
        assert response.status_code == 201


@pytest.mark.django_db
class TestFoodTicketViews:
    """Tests for food ticket API endpoints."""

    def test_list_tickets(self, authenticated_client, food_ticket):
        """Test listing available tickets."""
        url = reverse("food:ticket-list")
        response = authenticated_client.get(url)
        assert response.status_code == 200

    def test_create_ticket(self, api_client, user_with_house, monday_date):
        """Test creating a food ticket after deadline with an active registration."""
        api_client.force_authenticate(user=user_with_house)
        url = reverse("food:ticket-list")
        future_date = monday_date + timedelta(weeks=5)

        # Create an active registration so the ticket can be created
        MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=future_date,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        # Mock time to be after deadline (Thursday of the previous week)
        # Deadline is Wednesday 23:59:59 of the week before the meal
        mock_now_date = future_date - timedelta(days=4)  # Thursday before Monday
        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(datetime.combine(mock_now_date, time(10, 0)))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()
            data = {
                "date": future_date.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
                "price": "25.00",
            }
            response = api_client.post(url, data, format="json")
        assert response.status_code == 201

    def test_claim_ticket(self, api_client, admin_user, food_ticket):
        """Test claiming a food ticket."""
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:ticket-claim", kwargs={"pk": food_ticket.pk})
        with patch("apps.notifications.services.notify_ticket_claimed"):
            response = api_client.post(url)
        assert response.status_code == 200
        food_ticket.refresh_from_db()
        assert food_ticket.is_available is False
        assert food_ticket.claimed_by == admin_user

    def test_cannot_claim_own_house_ticket(self, api_client, user_with_house, food_ticket):
        """Test that user cannot claim (buy) their own house's ticket."""
        api_client.force_authenticate(user=user_with_house)
        url = reverse("food:ticket-claim", kwargs={"pk": food_ticket.pk})
        response = api_client.post(url)
        assert response.status_code == 400
        assert "eget hus" in response.json()["detail"].lower()
        food_ticket.refresh_from_db()
        assert food_ticket.is_available is True
        assert food_ticket.claimed_by is None

    def test_claim_ticket_sends_notification(self, api_client, admin_user, food_ticket):
        """Test that claiming another user's ticket sends notification."""
        api_client.force_authenticate(user=admin_user)
        with patch("apps.notifications.services.notify_ticket_claimed") as mock_notify:
            url = reverse("food:ticket-claim", kwargs={"pk": food_ticket.pk})
            response = api_client.post(url)
            assert response.status_code == 200
            # Notification should be called when claiming someone else's ticket
            mock_notify.assert_called_once()

    def test_cannot_claim_already_claimed_ticket(self, api_client, admin_user, food_ticket):
        """Test that already claimed tickets cannot be claimed again."""
        # First claim the ticket
        food_ticket.is_available = False
        food_ticket.claimed_by = admin_user
        food_ticket.save()

        # Try to claim again
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:ticket-claim", kwargs={"pk": food_ticket.pk})
        response = api_client.post(url)
        assert response.status_code == 400
        assert "no longer available" in response.json()["detail"].lower()

    def test_owner_cannot_claim_ticket_claimed_by_other(
        self, authenticated_client, food_ticket, admin_user
    ):
        """Test that owner cannot claim back a ticket that someone else claimed."""
        # Someone else claims the ticket
        food_ticket.is_available = False
        food_ticket.claimed_by = admin_user
        food_ticket.save()

        # Owner tries to claim it back
        url = reverse("food:ticket-claim", kwargs={"pk": food_ticket.pk})
        response = authenticated_client.post(url)
        assert response.status_code == 400
        assert "no longer available" in response.json()["detail"].lower()

    def test_after_claiming_own_ticket_can_register(self, api_client, user_with_house, monday_date):
        """Test that user can register even when they have an active ticket (partial selling)."""
        api_client.force_authenticate(user=user_with_house)

        # Create a ticket
        ticket = FoodTicket.objects.create(
            house=user_with_house.house,
            owner=user_with_house,
            date=monday_date,
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            price=Decimal("26.00"),
        )

        # Registration is NOT blocked by tickets - partial selling is allowed
        reg_url = reverse("food:registration-list")
        reg_data = {
            "date": monday_date.isoformat(),
            "adults_veg": 1,
            "children_count": 0,
            "is_active": True,
        }
        response = api_client.post(reg_url, reg_data, format="json")
        assert response.status_code == 201

        assert ticket is not None  # ticket still exists

    def test_release_ticket(self, api_client, user, admin_user, food_ticket):
        """Test releasing a claimed ticket."""
        # First claim the ticket
        food_ticket.is_available = False
        food_ticket.claimed_by = admin_user
        food_ticket.save()

        # Claimer releases it
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:ticket-release", kwargs={"pk": food_ticket.pk})
        response = api_client.post(url)
        assert response.status_code == 200
        food_ticket.refresh_from_db()
        assert food_ticket.is_available is True


@pytest.mark.django_db
class TestFoodTeamViews:
    """Tests for food team API endpoints."""

    def test_list_teams(self, authenticated_client, food_team):
        """Test listing food teams."""
        url = reverse("food:team-list")
        response = authenticated_client.get(url)
        assert response.status_code == 200

    def test_get_my_teams(self, authenticated_client, food_team, food_team_member):
        """Test getting user's teams."""
        url = reverse("food:my-teams")
        response = authenticated_client.get(url)
        assert response.status_code == 200
        data = response.json()
        if isinstance(data, dict) and "results" in data:
            data = data["results"]
        assert len(data) >= 1


@pytest.mark.django_db
class TestFoodTeamCycleViews:
    """Tests for food team cycle API endpoints."""

    def test_list_cycles(self, authenticated_client, food_team_cycle):
        """Test listing cycles."""
        url = reverse("food:cycle-list")
        response = authenticated_client.get(url)
        assert response.status_code == 200

    def test_create_cycle_requires_admin(self, authenticated_client, monday_date):
        """Test that only admins can create cycles."""
        url = reverse("food:cycle-list")
        cooking_dates = generate_cooking_dates(monday_date, num_weeks=1)
        data = {
            "name": "New Cycle",
            "cooking_dates": cooking_dates,
            "wish_deadline": (timezone.now() + timedelta(days=1)).isoformat(),
        }
        response = authenticated_client.post(url, data, format="json")
        assert response.status_code == 403

    def test_admin_can_create_cycle(self, admin_client, monday_date):
        """Test that admin can create cycles."""
        url = reverse("food:cycle-list")
        # Use a far future date to avoid conflicts
        far_monday = monday_date + timedelta(weeks=20)
        cooking_dates = generate_cooking_dates(far_monday, num_weeks=1)
        data = {
            "name": "Admin Cycle",
            "cooking_dates": cooking_dates,
            "wish_deadline": (timezone.now() + timedelta(days=1)).isoformat(),
        }
        response = admin_client.post(url, data, format="json")
        assert response.status_code == 201


@pytest.mark.django_db
class TestFoodTeamWishViews:
    """Tests for food team wish API endpoints."""

    def test_submit_wish(self, authenticated_client, food_team_cycle, monday_date):
        """Test submitting a wish."""
        url = reverse("food:my-wish", kwargs={"cycle_id": food_team_cycle.id})
        data = {
            "available_dates": [
                monday_date.isoformat(),
                (monday_date + timedelta(days=1)).isoformat(),
            ],
            "comment": "I prefer Mondays",
        }
        response = authenticated_client.post(url, data, format="json")
        assert response.status_code == 201

    def test_get_my_wish(self, authenticated_client, food_team_wish, food_team_cycle):
        """Test getting user's wish."""
        url = reverse("food:my-wish", kwargs={"cycle_id": food_team_cycle.id})
        response = authenticated_client.get(url)
        assert response.status_code == 200


@pytest.mark.django_db
class TestRegistrationStatsView:
    """Tests for registration statistics endpoint."""

    def test_get_daily_stats(self, authenticated_client, meal_registration, monday_date):
        """Test getting daily registration stats."""
        url = reverse("food:registration-stats")
        response = authenticated_client.get(url, {"date": monday_date.isoformat()})
        assert response.status_code == 200
        data = response.json()
        assert "total_registrations" in data
        assert "takeaway" in data
        assert "eat_in_1730" in data
        assert "eat_in_1830" in data

    def test_get_weekly_stats(self, authenticated_client, meal_registration, monday_date):
        """Test getting weekly registration stats."""
        url = reverse("food:registration-stats")
        response = authenticated_client.get(url, {"week_start": monday_date.isoformat()})
        assert response.status_code == 200


@pytest.mark.django_db
class TestDriveMenuViews:
    """Tests for drive menu API endpoints."""

    def test_get_drive_menu(self, authenticated_client, db):
        """Test getting drive menu (returns 404 if not cached)."""
        url = reverse("food:drive-menu")
        response = authenticated_client.get(url)
        # Will be 404 if no cache exists, 200 if it does
        assert response.status_code in [200, 404]

    def test_get_drive_menu_with_cache(self, authenticated_client, db):
        """Test getting drive menu when cache exists."""
        # Create cache entry for current week
        today = timezone.now().date()
        week_number = today.isocalendar()[1]
        year = today.isocalendar()[0]

        DriveMenuCache.objects.create(
            week_number=week_number,
            year=year,
            monday_menu="Lasagne",
            tuesday_menu="Thai curry",
            wednesday_menu="Frikadeller",
            thursday_menu="Pasta",
        )

        url = reverse("food:drive-menu")
        response = authenticated_client.get(url)
        assert response.status_code == 200
        data = response.json()
        assert data["monday_menu"] == "Lasagne"
        assert data["tuesday_menu"] == "Thai curry"

    def test_get_drive_menu_by_week(self, authenticated_client, db):
        """Test getting drive menu for specific week."""
        DriveMenuCache.objects.create(
            week_number=10,
            year=2026,
            monday_menu="Test menu",
        )

        url = reverse("food:drive-menu")
        response = authenticated_client.get(url, {"week": 10, "year": 2026})
        assert response.status_code == 200
        data = response.json()
        assert data["week_number"] == 10
        assert data["year"] == 2026

    def test_refresh_drive_menu_requires_admin(self, authenticated_client):
        """Test that refreshing drive menu requires admin."""
        url = reverse("food:drive-menu")
        response = authenticated_client.post(url)
        assert response.status_code == 403


# =============================================================================
# Team Generator Service Tests
# =============================================================================


@pytest.mark.django_db
class TestTeamGenerator:
    """Tests for TeamGenerator service."""

    def test_generator_initialization(self, food_team_cycle):
        """Test generator initialization."""
        generator = TeamGenerator(food_team_cycle)
        assert generator.cycle == food_team_cycle
        assert len(generator.cooking_dates) > 0

    def test_load_data(self, food_team_cycle, user):
        """Test loading data from database."""
        # Mark user as eligible
        user.is_exempt_from_food_teams = False
        user.save()

        generator = TeamGenerator(food_team_cycle)
        generator.load_data()

        assert len(generator.persons) >= 1

    def test_is_valid_assignment(self, food_team_cycle, multiple_users, monday_date):
        """Test assignment validation."""
        generator = TeamGenerator(food_team_cycle)
        generator.load_data()

        # First assignment should be valid
        user_id = multiple_users[0].id
        if user_id in generator.persons and monday_date in generator.cooking_dates:
            assert generator.is_valid_assignment(user_id, monday_date) is True

    def test_generate_dry_run(self, food_team_cycle, multiple_users, monday_date):
        """Test dry run generation."""
        # Create wishes for all users (cooking_dates is already a list of ISO strings)
        for user in multiple_users:
            user.is_exempt_from_food_teams = False
            user.save()
            FoodTeamWish.objects.create(
                cycle=food_team_cycle,
                user=user,
                available_dates=food_team_cycle.cooking_dates[:8],
            )

        generator = TeamGenerator(food_team_cycle)
        result = generator.generate(save=False)

        assert isinstance(result, TeamGenerationResult)
        # Should not have saved teams
        assert FoodTeam.objects.filter(cycle=food_team_cycle).count() == 0

    def test_generate_with_save(self, db, admin_user, multiple_users):
        """Test actual generation with save."""
        # Create a fresh cycle
        today = timezone.now().date()
        next_monday = today + timedelta(days=(7 - today.weekday()))
        cooking_dates = generate_cooking_dates(next_monday, num_weeks=2)

        cycle = FoodTeamCycle.objects.create(
            name="Generation Test Cycle",
            cooking_dates=cooking_dates,
            wish_deadline=timezone.now() + timedelta(days=7),
            status=CycleStatus.COLLECTING_WISHES,
            created_by=admin_user,
        )

        # Create wishes (cooking_dates is already a list of ISO strings)
        for user in multiple_users:
            user.is_exempt_from_food_teams = False
            user.save()
            FoodTeamWish.objects.create(
                cycle=cycle,
                user=user,
                available_dates=cycle.cooking_dates,
            )

        generator = TeamGenerator(cycle)
        result = generator.generate(save=True)

        assert isinstance(result, TeamGenerationResult)
        assert result.teams_created > 0

        # Verify teams were saved
        teams = FoodTeam.objects.filter(cycle=cycle)
        assert teams.count() > 0

    def test_house_conflict_detection(self, db, admin_user, house, monday_date):
        """Test that users from same house are not assigned together."""
        # Create users from same house
        users = []
        for i in range(3):
            user = User.objects.create_user(
                email=f"sameHouse{i}@example.com",
                password="testpass123",
                first_name=f"SameHouse{i}",
                house=house,
            )
            users.append(user)

        # Create cycle
        cooking_dates = generate_cooking_dates(monday_date, num_weeks=1)
        cycle = FoodTeamCycle.objects.create(
            name="House Conflict Test",
            cooking_dates=cooking_dates,
            wish_deadline=timezone.now() + timedelta(days=7),
            status=CycleStatus.COLLECTING_WISHES,
            created_by=admin_user,
        )

        generator = TeamGenerator(cycle)
        generator.load_data()

        # Assign first user (cooking_dates are now date objects in the generator)
        monday_date_iso = monday_date.isoformat()
        if monday_date_iso in cycle.cooking_dates and users[0].id in generator.persons:
            generator.assign_person(users[0].id, monday_date)

            # Second user from same house should be invalid
            if users[1].id in generator.persons:
                assert generator.is_valid_assignment(users[1].id, monday_date) is False

    def test_over_50_limit(self, db, admin_user, house, house2, monday_date):
        """Test that max 2 over-50 people are assigned per team."""
        # Create over-50 users
        users = []
        for i in range(4):
            user = User.objects.create_user(
                email=f"over50User{i}@example.com",
                password="testpass123",
                first_name=f"Over50User{i}",
                house=house if i % 2 == 0 else house2,  # Different houses
                is_over_50=True,
            )
            users.append(user)

        cooking_dates = generate_cooking_dates(monday_date, num_weeks=1)
        cycle = FoodTeamCycle.objects.create(
            name="Over 50 Test",
            cooking_dates=cooking_dates,
            wish_deadline=timezone.now() + timedelta(days=7),
            status=CycleStatus.COLLECTING_WISHES,
            created_by=admin_user,
        )

        generator = TeamGenerator(cycle)
        generator.load_data()

        if monday_date in generator.cooking_dates:
            # Assign first two over-50 users
            for i in range(2):
                if users[i].id in generator.persons:
                    generator.assign_person(users[i].id, monday_date)

            # Third over-50 user should be invalid
            if users[2].id in generator.persons:
                assert generator.is_valid_assignment(users[2].id, monday_date) is False


# =============================================================================
# Apply Defaults View Tests
# =============================================================================


@pytest.mark.django_db
class TestWednesdayMeatVeg:
    """Tests for Wednesday mixed meal types."""

    def test_wednesday_mixed_registration(self, api_client, user_with_house, monday_date):
        """Test registering mixed meat+veg on Wednesday."""
        api_client.force_authenticate(user=user_with_house)
        wednesday = monday_date + timedelta(days=2)
        url = reverse("food:registration-list")
        data = {
            "date": wednesday.isoformat(),
            "adults_meat": 1,
            "adults_veg": 1,
            "children_count": 0,
        }
        response = api_client.post(url, data, format="json")
        assert response.status_code == 201
        result = response.json()
        assert result["adults_meat"] == 1
        assert result["adults_veg"] == 1

    def test_wednesday_preference_allows_meat(self, api_client, user_with_house):
        """Test that a Wednesday preference can have adults_meat > 0."""
        api_client.force_authenticate(user=user_with_house)
        url = reverse("food:preference-list")
        data = {
            "day_of_week": 2,  # Wednesday
            "adults_meat": 2,
            "adults_veg": 0,
            "children_count": 0,
        }
        response = api_client.post(url, data, format="json")
        assert response.status_code == 201

    def test_non_wednesday_preference_rejects_meat(self, api_client, user_with_house):
        """Test that a non-Wednesday preference rejects adults_meat > 0."""
        api_client.force_authenticate(user=user_with_house)
        url = reverse("food:preference-list")
        data = {
            "day_of_week": 1,  # Tuesday
            "adults_meat": 1,
            "adults_veg": 0,
            "children_count": 0,
        }
        response = api_client.post(url, data, format="json")
        assert response.status_code == 400

    def test_stats_include_meat_veg_breakdown(self, authenticated_client, monday_date, house):
        """Stats for total include adults_meat and adults_veg breakdown."""
        user = (
            authenticated_client.handler._force_user
            if hasattr(authenticated_client, "handler")
            else User.objects.get(email="test@example.com")
        )
        user.house = house
        user.save()
        wednesday = monday_date + timedelta(days=2)
        MealRegistration.objects.create(
            house=house,
            last_modified_by=user,
            date=wednesday,
            adults_meat=2,
            adults_veg=1,
            children_count=0,
            is_active=True,
        )
        url = reverse("food:registration-stats")
        response = authenticated_client.get(url, {"date": wednesday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        assert "adults_meat" in data["total_registrations"]
        assert "adults_veg" in data["total_registrations"]


@pytest.mark.django_db
class TestVirtualRegistrations:
    """Tests for virtual (preference-based) registrations."""

    def test_week_start_returns_4_rows(self, api_client, user_with_house, monday_date):
        """GET with week_start always returns exactly 4 rows (Mon-Thu)."""
        api_client.force_authenticate(user=user_with_house)
        future_monday = monday_date + timedelta(weeks=5)
        url = reverse("food:registration-list")
        response = api_client.get(url, {"week_start": future_monday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        assert len(data) == 4

    def test_virtual_rows_have_null_id(self, api_client, user_with_house, monday_date):
        """Days without DB rows have id=null and is_from_preference=True."""
        api_client.force_authenticate(user=user_with_house)
        future_monday = monday_date + timedelta(weeks=5)
        url = reverse("food:registration-list")
        response = api_client.get(url, {"week_start": future_monday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        for row in data:
            assert row["id"] is None
            assert row["is_from_preference"] is True

    def test_real_rows_have_real_ids(self, api_client, user_with_house, monday_date):
        """Days with DB rows have real ids and is_from_preference=False."""
        api_client.force_authenticate(user=user_with_house)
        future_monday = monday_date + timedelta(weeks=5)
        MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=future_monday,
            adults_veg=2,
            is_active=True,
        )
        url = reverse("food:registration-list")
        response = api_client.get(url, {"week_start": future_monday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        monday_row = next(r for r in data if r["date"] == future_monday.isoformat())
        assert monday_row["id"] is not None
        assert monday_row["is_from_preference"] is False
        # Other days are still virtual
        other_rows = [r for r in data if r["date"] != future_monday.isoformat()]
        for row in other_rows:
            assert row["id"] is None
            assert row["is_from_preference"] is True

    def test_virtual_values_match_preferences(self, api_client, user_with_house, monday_date):
        """Virtual registration values match user preferences."""
        api_client.force_authenticate(user=user_with_house)
        MealPreference.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            day_of_week=0,  # Monday
            adults_veg=3,
            adults_meat=0,
            children_count=1,
            dining_option="take_away",
            seating_time="17:30",
        )
        future_monday = monday_date + timedelta(weeks=5)
        url = reverse("food:registration-list")
        response = api_client.get(url, {"week_start": future_monday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        monday_row = next(r for r in data if r["date"] == future_monday.isoformat())
        assert monday_row["adults_veg"] == 3
        assert monday_row["children_count"] == 1
        assert monday_row["dining_option"] == "take_away"

    def test_virtual_defaults_when_no_preference(self, api_client, user_with_house, monday_date):
        """Virtual registration uses house_count veg portions when no preference."""
        api_client.force_authenticate(user=user_with_house)
        house_count = user_with_house.house.inhabitants.count()
        future_monday = monday_date + timedelta(weeks=5)
        url = reverse("food:registration-list")
        response = api_client.get(url, {"week_start": future_monday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        monday_row = next(r for r in data if r["date"] == future_monday.isoformat())
        assert monday_row["adults_veg"] == house_count
        assert monday_row["adults_meat"] == 0
        assert monday_row["children_count"] == 0


@pytest.mark.django_db
class TestLazyMaterialization:
    """Tests for lazy materialization of registrations post-deadline."""

    def test_post_deadline_creates_real_row(self, api_client, user_with_house):
        """GET for a post-deadline week creates real DB rows."""
        api_client.force_authenticate(user=user_with_house)
        # Use a hardcoded past Monday (deadline definitely passed)
        past_monday = date(2024, 1, 8)  # A Monday far in the past

        url = reverse("food:registration-list")
        response = api_client.get(url, {"week_start": past_monday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        assert len(data) == 4
        # All rows should have real ids (materialized)
        for row in data:
            assert row["id"] is not None
            assert row["is_from_preference"] is False

        # Real DB rows were created
        assert MealRegistration.objects.filter(
            house=user_with_house.house, date=past_monday
        ).exists()

    def test_materialized_values_match_preferences(self, api_client, user_with_house):
        """Materialized registrations use preference values."""
        api_client.force_authenticate(user=user_with_house)
        MealPreference.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            day_of_week=0,  # Monday
            adults_veg=3,
            adults_meat=0,
            children_count=0,
            dining_option="eat_in",
            seating_time="17:30",
        )
        past_monday = date(2024, 2, 5)  # A Monday far in the past

        url = reverse("food:registration-list")
        response = api_client.get(url, {"week_start": past_monday.isoformat()})
        assert response.status_code == 200
        data = response.json()
        monday_row = next(r for r in data if r["date"] == past_monday.isoformat())
        assert monday_row["adults_veg"] == 3

        # Verify DB row matches
        reg = MealRegistration.objects.get(house=user_with_house.house, date=past_monday)
        assert reg.adults_veg == 3

    def test_concurrent_materialization_no_error(self, api_client, user_with_house):
        """Concurrent materialization uses get_or_create — no integrity errors."""
        from apps.food.views import _materialize_registration

        past_monday = date(2024, 3, 4)  # A Monday far in the past

        # Call twice (simulates concurrent requests)
        reg1 = _materialize_registration(
            user_with_house, past_monday, None, user_with_house.house, 1
        )
        reg2 = _materialize_registration(
            user_with_house, past_monday, None, user_with_house.house, 1
        )
        assert reg1.id == reg2.id


@pytest.mark.django_db
class TestStatsWithPreferenceFallback:
    """Tests for stats endpoint including preference-based virtual contributions."""

    def test_stats_nonzero_with_preferences_no_real_regs(
        self, api_client, user_with_house, monday_date
    ):
        """Stats for a date with no real registrations is non-zero when preferences exist."""
        api_client.force_authenticate(user=user_with_house)
        MealPreference.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            day_of_week=monday_date.weekday(),
            adults_veg=2,
            adults_meat=0,
            children_count=0,
            dining_option="eat_in",
            seating_time="17:30",
        )
        url = reverse("food:registration-stats")
        response = api_client.get(url, {"date": monday_date.isoformat()})
        assert response.status_code == 200
        data = response.json()
        assert data["total_registrations"]["adults_veg"] > 0

    def test_user_with_real_reg_not_double_counted(self, api_client, user_with_house, monday_date):
        """House with a real registration is not also counted via preference."""
        api_client.force_authenticate(user=user_with_house)
        MealPreference.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            day_of_week=monday_date.weekday(),
            adults_veg=3,
            adults_meat=0,
            children_count=0,
            dining_option="eat_in",
            seating_time="17:30",
        )
        MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=monday_date,
            adults_veg=2,
            is_active=True,
        )
        url = reverse("food:registration-stats")
        response = api_client.get(url, {"date": monday_date.isoformat()})
        assert response.status_code == 200
        data = response.json()
        # Should count 2 (real reg), not 5 (2 + 3)
        assert data["total_registrations"]["adults_veg"] == 2

    def test_two_users_same_house_counted_once(self, api_client, house, monday_date):
        """House with a preference is counted only once regardless of inhabitants."""
        user1 = User.objects.create_user(
            email="housemate1@example.com", password="pass", first_name="A", house=house
        )
        User.objects.create_user(
            email="housemate2@example.com", password="pass", first_name="B", house=house
        )
        api_client.force_authenticate(user=user1)
        # One preference per house (unique constraint: house + day_of_week)
        MealPreference.objects.create(
            house=house,
            last_modified_by=user1,
            day_of_week=monday_date.weekday(),
            adults_veg=2,
            adults_meat=0,
            children_count=0,
            dining_option="eat_in",
            seating_time="17:30",
        )
        url = reverse("food:registration-stats")
        response = api_client.get(url, {"date": monday_date.isoformat()})
        assert response.status_code == 200
        data = response.json()
        # Only one preference per house
        assert data["total_registrations"]["adults_veg"] == 2


@pytest.mark.django_db
class TestBillingDedup:
    """Tests for billing deduplication (fixes double-counting bug)."""

    def test_one_registration_per_house_billed_once(self, api_client, admin_user, house):
        """One registration per house per date is billed once (unique constraint: house+date)."""
        # Use a hardcoded past date so deadline has passed
        past_monday = date(2024, 4, 1)  # Monday, April 2024

        user1 = User.objects.create_user(
            email="bill1@example.com", password="pass", first_name="Bill1", house=house
        )
        # Zero-portion preferences so other Mon-Thu dates don't get default billing
        for day in range(4):
            MealPreference.objects.get_or_create(
                house=house,
                day_of_week=day,
                defaults={
                    "adults_meat": 0,
                    "adults_veg": 0,
                    "children_count": 0,
                    "dining_option": "eat_in",
                    "seating_time": "17:30",
                },
            )
        MealRegistration.objects.create(
            house=house, last_modified_by=user1, date=past_monday, adults_veg=2, is_active=True
        )

        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        first = past_monday.replace(day=1)
        import calendar as _cal

        last = past_monday.replace(day=_cal.monthrange(past_monday.year, past_monday.month)[1])
        response = api_client.get(
            url, {"start_date": first.isoformat(), "end_date": last.isoformat()}
        )
        assert response.status_code == 200
        data = response.json()
        house_data = next(h for h in data["houses"] if h["house_id"] == house.id)
        # Should count 2 portions (one registration per house)
        assert house_data["adult_veg_portions"] == 2


@pytest.mark.django_db
class TestCostReportPermissions:
    """The food cost report is visible to food admins and economy admins."""

    def test_economy_admin_can_view_cost_report(self, api_client, house):
        econ = User.objects.create_user(
            email="econ@example.com",
            password="pass12345",
            first_name="Econ",
            is_economy_admin=True,
        )
        api_client.force_authenticate(user=econ)
        url = reverse("food:monthly-food-cost")
        resp = api_client.get(url, {"start_date": "2024-04-01", "end_date": "2024-04-30"})
        assert resp.status_code == 200

    def test_regular_user_cannot_view_cost_report(self, api_client):
        plain = User.objects.create_user(
            email="plain@example.com", password="pass12345", first_name="Plain"
        )
        api_client.force_authenticate(user=plain)
        url = reverse("food:monthly-food-cost")
        resp = api_client.get(url, {"start_date": "2024-04-01", "end_date": "2024-04-30"})
        assert resp.status_code == 403


@pytest.mark.django_db
class TestBillingWithPreferenceFallback:
    """Tests for billing using preference/default fallback."""

    def test_billing_includes_house_with_preference_no_reg(self, api_client, admin_user, house):
        """A house with a preference but no registration is still billed."""
        user1 = User.objects.create_user(
            email="prefbill@example.com", password="pass", first_name="PB", house=house
        )
        MealPreference.objects.create(
            house=house,
            last_modified_by=user1,
            day_of_week=0,  # Monday
            adults_veg=3,
            adults_meat=0,
            children_count=0,
            dining_option="eat_in",
            seating_time="17:30",
        )
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2024-04-01", "end_date": "2024-04-30"})
        assert response.status_code == 200
        data = response.json()
        house_data = next(h for h in data["houses"] if h["house_id"] == house.id)
        assert house_data["adult_veg_portions"] > 0

    def test_billing_uses_default_for_house_with_no_preference(self, api_client, admin_user, house):
        """A house with inhabitants but no preference is billed house_count veg."""
        User.objects.create_user(
            email="nopref@example.com", password="pass", first_name="NP", house=house
        )
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        # Use a past month so all Mon-Thu dates have passed the deadline
        response = api_client.get(url, {"start_date": "2024-04-01", "end_date": "2024-04-30"})
        assert response.status_code == 200
        data = response.json()
        house_data = next(h for h in data["houses"] if h["house_id"] == house.id)
        # House has at least 1 inhabitant, so should be billed for each Mon-Thu
        assert house_data["adult_veg_portions"] > 0

    def test_billing_inactive_reg_not_billed_via_fallback(self, api_client, admin_user, house):
        """An explicitly cancelled registration (is_active=False) is not billed."""
        past_monday = date(2024, 5, 6)  # Monday, May 2024
        user1 = User.objects.create_user(
            email="inactive@example.com", password="pass", first_name="IA", house=house
        )
        # Set zero preferences so only the explicit registration matters
        for day in range(4):
            MealPreference.objects.create(
                house=house,
                last_modified_by=user1,
                day_of_week=day,
                adults_meat=0,
                adults_veg=0,
                children_count=0,
                dining_option="eat_in",
                seating_time="17:30",
            )
        # Explicit opt-out via is_active=False registration
        MealRegistration.objects.create(
            house=house, last_modified_by=user1, date=past_monday, adults_veg=3, is_active=False
        )
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2024-05-01", "end_date": "2024-05-31"})
        assert response.status_code == 200
        data = response.json()
        house_data = next(h for h in data["houses"] if h["house_id"] == house.id)
        # Cancelled registration should not be billed
        assert house_data["total_cost"] == "0.00"


# =============================================================================
# Unauthenticated Access Tests
# =============================================================================


@pytest.mark.django_db
class TestUnauthenticatedAccess:
    """Tests to ensure endpoints require authentication."""

    def test_ticket_list_requires_auth(self, api_client):
        """Test that ticket list requires authentication."""
        url = reverse("food:ticket-list")
        response = api_client.get(url)
        assert response.status_code == 401

    def test_team_list_requires_auth(self, api_client):
        """Test that team list requires authentication."""
        url = reverse("food:team-list")
        response = api_client.get(url)
        assert response.status_code == 401

    def test_drive_menu_requires_auth(self, api_client):
        """Test that drive menu requires authentication."""
        url = reverse("food:drive-menu")
        response = api_client.get(url)
        assert response.status_code == 401


# =============================================================================
# Food Ticket Default Pricing and Registration Tests
# =============================================================================


@pytest.mark.django_db
class TestFoodTicketDefaultPricing:
    """Tests for food ticket default pricing."""

    def test_default_price_meat(self, db, user):
        """Test default price calculation for meat meal."""
        serializer = FoodTicketCreateSerializer(context={"request": MockRequest(user)})
        price = serializer.calculate_default_price(
            adults_meat=2, adults_veg=0, children_count=1, meal_date=date(2025, 12, 22)
        )
        # 2 adults @ 37 + 1 child @ 18 = 92
        assert price == Decimal("92.00")

    def test_default_price_vegetarian(self, db, user):
        """Test default price calculation for vegetarian meal."""
        serializer = FoodTicketCreateSerializer(context={"request": MockRequest(user)})
        price = serializer.calculate_default_price(
            adults_meat=0, adults_veg=2, children_count=1, meal_date=date(2025, 12, 22)
        )
        # 2 adults @ 26 + 1 child @ 18 = 70
        assert price == Decimal("70.00")

    def test_default_price_mixed(self, db, user):
        """Test default price calculation for mixed meat+veg meal."""
        serializer = FoodTicketCreateSerializer(context={"request": MockRequest(user)})
        price = serializer.calculate_default_price(
            adults_meat=1, adults_veg=1, children_count=1, meal_date=date(2025, 12, 22)
        )
        # 1 adult @ 37 + 1 adult @ 26 + 1 child @ 18 = 81
        assert price == Decimal("81.00")

    def test_default_price_uses_meal_date_not_today(self, db, user):
        """A meal on/after the 2026-08-02 increase is priced at the new rates."""
        serializer = FoodTicketCreateSerializer(context={"request": MockRequest(user)})
        price = serializer.calculate_default_price(
            adults_meat=1, adults_veg=1, children_count=1, meal_date=date(2026, 8, 3)
        )
        # 1 adult @ 40 + 1 adult @ 30 + 1 child @ 18 = 88
        assert price == Decimal("88.00")

    def test_ticket_created_with_default_price(self, api_client, user_with_house, monday_date):
        """Test that ticket is created with default price if not specified."""
        api_client.force_authenticate(user=user_with_house)

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(
                timezone.datetime(2025, 12, 18, 10, 0)  # Thursday
            )
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            # Ticket for next Monday (veg - no meat on Monday)
            next_monday = date(2025, 12, 22)

            # Create active registration first (required to sell a ticket)
            MealRegistration.objects.create(
                house=user_with_house.house,
                last_modified_by=user_with_house,
                date=next_monday,
                adults_meat=0,
                adults_veg=2,
                children_count=1,
                is_active=True,
            )

            url = reverse("food:ticket-list")
            data = {
                "date": next_monday.isoformat(),
                "adults_meat": 0,
                "adults_veg": 2,
                "children_count": 1,
                # No price specified - should use default
            }
            response = api_client.post(url, data, format="json")

            assert response.status_code == 201
            # Default price: 2 adults @ 26 + 1 child @ 18 = 70
            assert Decimal(response.json()["price"]) == Decimal("70.00")


@pytest.mark.django_db
class TestPartialTicketSelling:
    """Tests for partial ticket selling.

    In the new model, tickets do NOT modify the registration.
    The registration is the immutable billing record.
    available_portions = registration portions - sum of existing available tickets.
    """

    def test_ticket_does_not_modify_registration(self, api_client, user_with_house, monday_date):
        """Creating a ticket (even for all portions) does NOT modify the registration."""
        api_client.force_authenticate(user=user_with_house)

        reg_date = date(2025, 12, 22)  # A Monday
        registration = MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=reg_date,
            adults_meat=0,
            adults_veg=2,
            children_count=1,
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:ticket-list")
            data = {
                "date": reg_date.isoformat(),
                "adults_meat": 0,
                "adults_veg": 2,
                "children_count": 1,
            }
            response = api_client.post(url, data, format="json")
            assert response.status_code == 201

        # Registration is UNCHANGED
        registration.refresh_from_db()
        assert registration.is_active is True
        assert registration.adults_veg == 2
        assert registration.children_count == 1

    def test_partial_ticket_registration_unchanged(self, api_client, user_with_house, monday_date):
        """Selling a partial portion leaves the registration unchanged."""
        api_client.force_authenticate(user=user_with_house)

        reg_date = date(2025, 12, 22)  # A Monday
        registration = MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=reg_date,
            adults_meat=0,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:ticket-list")
            data = {
                "date": reg_date.isoformat(),
                "adults_meat": 0,
                "adults_veg": 1,  # Sell 1 of 2
                "children_count": 0,
            }
            response = api_client.post(url, data, format="json")
            assert response.status_code == 201

        # Registration is UNCHANGED
        registration.refresh_from_db()
        assert registration.is_active is True
        assert registration.adults_veg == 2  # still 2, not reduced

    def test_overselling_rejected(self, api_client, user_with_house, monday_date):
        """Cannot sell more portions than registered."""
        api_client.force_authenticate(user=user_with_house)

        reg_date = date(2025, 12, 22)  # A Monday
        MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=reg_date,
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:ticket-list")
            data = {
                "date": reg_date.isoformat(),
                "adults_meat": 0,
                "adults_veg": 3,  # More than registered
                "children_count": 0,
            }
            response = api_client.post(url, data, format="json")
            assert response.status_code == 400

    def test_ticket_deletion_does_not_modify_registration(
        self, api_client, user_with_house, monday_date
    ):
        """Deleting a ticket does NOT modify the registration (registrations are immutable)."""
        api_client.force_authenticate(user=user_with_house)

        reg_date = date(2025, 12, 22)  # A Monday
        registration = MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=reg_date,
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            is_active=True,
        )

        ticket = FoodTicket.objects.create(
            house=user_with_house.house,
            owner=user_with_house,
            date=reg_date,
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            is_available=True,
        )

        url = reverse("food:ticket-detail", kwargs={"pk": ticket.pk})
        response = api_client.delete(url)
        assert response.status_code == 204

        # Registration is UNCHANGED — deleting a ticket does not restore portions
        registration.refresh_from_db()
        assert registration.is_active is True
        assert registration.adults_veg == 1  # unchanged

    def test_registration_coexists_with_ticket(self, api_client, user_with_house, monday_date):
        """User can have both a registration and a ticket (partial selling)."""
        api_client.force_authenticate(user=user_with_house)

        ticket_date = monday_date + timedelta(weeks=4)

        # Create an available ticket - registration is NOT blocked
        FoodTicket.objects.create(
            house=user_with_house.house,
            owner=user_with_house,
            date=ticket_date,
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            is_available=True,
        )

        # User can still register for the same date
        url = reverse("food:registration-list")
        data = {
            "date": ticket_date.isoformat(),
            "adults_veg": 1,
            "children_count": 0,
        }
        response = api_client.post(url, data, format="json")
        assert response.status_code == 201


@pytest.mark.django_db
class TestTicketCreationDeadline:
    """Tests for ticket creation deadline restrictions."""

    def test_cannot_create_ticket_before_deadline(self, api_client, user_with_house):
        """Test that tickets cannot be created before registration deadline."""
        api_client.force_authenticate(user=user_with_house)

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 15, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            next_monday = date(2025, 12, 22)

            url = reverse("food:ticket-list")
            data = {
                "date": next_monday.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
            }
            response = api_client.post(url, data, format="json")

            assert response.status_code == 400
            assert "frist" in str(response.json()).lower()

    def test_can_create_ticket_after_deadline(self, api_client, user_with_house):
        """Test that tickets can be created after registration deadline (with active reg)."""
        api_client.force_authenticate(user=user_with_house)

        next_monday = date(2025, 12, 22)

        # Create an active registration first
        MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=next_monday,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:ticket-list")
            data = {
                "date": next_monday.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
            }
            response = api_client.post(url, data, format="json")

            assert response.status_code == 201

    def test_deadline_calculation(self, db, user):
        """Test deadline calculation for different meal dates."""
        from apps.food.serializers import get_registration_deadline

        # Monday Dec 22, 2025 - deadline should be Wed Dec 17 at 23:59:59
        monday = date(2025, 12, 22)
        deadline = get_registration_deadline(monday)
        assert deadline.date() == date(2025, 12, 17)
        assert deadline.hour == 23
        assert deadline.minute == 59

        # Tuesday Dec 23 - deadline should also be Wed Dec 17 at 23:59:59
        tuesday = date(2025, 12, 23)
        deadline = get_registration_deadline(tuesday)
        assert deadline.date() == date(2025, 12, 17)

        # Wednesday Dec 24 - deadline should be Wed Dec 17 at 23:59:59
        wednesday = date(2025, 12, 24)
        deadline = get_registration_deadline(wednesday)
        assert deadline.date() == date(2025, 12, 17)

        # Thursday Dec 25 - deadline should be Wed Dec 17 at 23:59:59
        thursday = date(2025, 12, 25)
        deadline = get_registration_deadline(thursday)
        assert deadline.date() == date(2025, 12, 17)


@pytest.mark.django_db
class TestMonthlyFoodCostReport:
    """Tests for monthly food cost report.

    Billing is based on active meal registrations + preference/default fallback.
    Ticket sales are between users (MobilePay) and do NOT affect the house bill.
    """

    @staticmethod
    def _set_zero_preferences(user):
        """Set zero-portion preferences for all days, opting the house out of defaults."""
        for day in range(4):
            MealPreference.objects.get_or_create(
                house=user.house,
                day_of_week=day,
                defaults={
                    "adults_meat": 0,
                    "adults_veg": 0,
                    "children_count": 0,
                    "dining_option": "eat_in",
                    "seating_time": "17:30",
                    "last_modified_by": user,
                },
            )

    def test_cost_charged_to_registration_house(
        self, api_client, admin_user, user_with_house, house
    ):
        """Test that food cost is charged to the registration owner's house."""
        # Clean up existing data for this month
        FoodTicket.objects.filter(date__year=2025, date__month=1).delete()
        MealRegistration.objects.filter(date__year=2025, date__month=1).delete()
        self._set_zero_preferences(user_with_house)

        reg_date = date(2025, 1, 13)  # A Monday in January 2025
        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=reg_date,
            adults_meat=0,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        # Get the monthly cost report as admin
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2025-01-01", "end_date": "2025-01-31"})

        assert response.status_code == 200
        data = response.json()

        owner_house_cost = next((h for h in data["houses"] if h["house_id"] == house.id), None)

        # Cost: 2 adults * 26 DKK (veg) = 52 DKK
        assert owner_house_cost is not None
        assert Decimal(owner_house_cost["total_cost"]) == Decimal("52.00")
        assert owner_house_cost["adult_meat_portions"] >= 0
        assert owner_house_cost["registration_count"] == 1

    def test_only_admin_can_access_report(self, api_client, user_with_house):
        """Test that non-admin users cannot access the cost report."""
        api_client.force_authenticate(user=user_with_house)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2025-01-01", "end_date": "2025-01-31"})

        assert response.status_code == 403

    def test_report_defaults_and_validation(self, api_client, admin_user):
        """No params → default range (last 4 weeks); bad inputs → 400."""
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")

        # No params: returns default range (last 4 weeks)
        response = api_client.get(url)
        assert response.status_code == 200
        body = response.json()
        assert "start_date" in body and "end_date" in body

        # Invalid date format
        response = api_client.get(url, {"start_date": "not-a-date"})
        assert response.status_code == 400

        # end_date before start_date
        response = api_client.get(url, {"start_date": "2025-02-01", "end_date": "2025-01-01"})
        assert response.status_code == 400

    def test_report_totals_multiple_registrations(
        self, api_client, admin_user, user_with_house, house
    ):
        """Test that report correctly sums multiple registrations."""
        # Clean up existing data for this test
        FoodTicket.objects.filter(date__year=2025, date__month=2).delete()
        MealRegistration.objects.filter(date__year=2025, date__month=2).delete()
        self._set_zero_preferences(user_with_house)

        # Create multiple active registrations (non-Wednesday → veg)
        for i in range(3):
            MealRegistration.objects.create(
                house=house,
                last_modified_by=user_with_house,
                date=date(2025, 2, 3 + i),  # Feb 3, 4, 5 (Mon, Tue, Wed)
                adults_meat=0,
                adults_veg=1,
                children_count=0,
                is_active=True,
            )

        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2025-02-01", "end_date": "2025-02-28"})

        assert response.status_code == 200
        data = response.json()

        owner_house_cost = next((h for h in data["houses"] if h["house_id"] == house.id), None)

        assert owner_house_cost is not None
        assert Decimal(owner_house_cost["total_cost"]) == Decimal("78.00")  # 3 * 26 (veg)
        assert owner_house_cost["adult_meat_portions"] >= 0
        assert owner_house_cost["registration_count"] == 3

    def test_tickets_not_billed(self, api_client, admin_user, user_with_house, house):
        """Test that tickets are NOT charged to the owner's house (billing is registration-based)."""
        # Clean up existing data for this test
        FoodTicket.objects.filter(date__year=2025, date__month=3).delete()
        MealRegistration.objects.filter(date__year=2025, date__month=3).delete()
        self._set_zero_preferences(user_with_house)

        # Create a ticket — this should NOT affect the house bill
        FoodTicket.objects.create(
            house=user_with_house.house,
            owner=user_with_house,
            date=date(2025, 3, 3),  # A Monday in March 2025
            adults_meat=0,
            adults_veg=1,
            children_count=1,
            price=Decimal("44.00"),
            is_available=True,
        )

        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2025-03-01", "end_date": "2025-03-31"})

        assert response.status_code == 200
        data = response.json()

        owner_house_cost = next((h for h in data["houses"] if h["house_id"] == house.id), None)

        # Tickets do NOT contribute to the house bill
        assert owner_house_cost is not None
        assert Decimal(owner_house_cost["total_cost"]) == Decimal("0.00")
        assert owner_house_cost["adult_meat_portions"] >= 0
        assert owner_house_cost["registration_count"] == 0

    def test_active_registrations_charged(self, api_client, admin_user, user_with_house, house):
        """Test that active meal registrations are included in the cost report."""
        # Clean up existing data for this test
        FoodTicket.objects.filter(date__year=2025, date__month=4).delete()
        MealRegistration.objects.filter(date__year=2025, date__month=4).delete()
        self._set_zero_preferences(user_with_house)

        # Create an active meal registration (user ate the meal) - Monday is veg
        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=date(2025, 4, 7),  # A Monday in April 2025
            adults_meat=0,
            adults_veg=2,
            children_count=1,
            is_active=True,
        )

        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2025-04-01", "end_date": "2025-04-30"})

        assert response.status_code == 200
        data = response.json()

        owner_house_cost = next((h for h in data["houses"] if h["house_id"] == house.id), None)

        # Cost should be: 2 adults * 26 (veg) + 1 child * 18 = 70 DKK
        assert owner_house_cost is not None
        assert Decimal(owner_house_cost["total_cost"]) == Decimal("70.00")
        assert owner_house_cost["registration_count"] == 1
        assert owner_house_cost["adult_meat_portions"] >= 0

    def test_inactive_registrations_not_charged(
        self, api_client, admin_user, user_with_house, house
    ):
        """Test that inactive meal registrations (cancelled) are NOT charged."""
        # Clean up existing data for this test
        FoodTicket.objects.filter(date__year=2025, date__month=5).delete()
        MealRegistration.objects.filter(date__year=2025, date__month=5).delete()
        self._set_zero_preferences(user_with_house)

        # Create an inactive meal registration (user cancelled without creating ticket)
        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=date(2025, 5, 5),  # A Monday in May 2025
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            is_active=False,  # Cancelled
        )

        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2025-05-01", "end_date": "2025-05-31"})

        assert response.status_code == 200
        data = response.json()

        owner_house_cost = next((h for h in data["houses"] if h["house_id"] == house.id), None)

        # No cost - the registration was cancelled and no ticket was created
        assert owner_house_cost is not None
        assert Decimal(owner_house_cost["total_cost"]) == Decimal("0.00")
        assert owner_house_cost["registration_count"] == 0
        assert owner_house_cost["adult_meat_portions"] >= 0

    def test_only_active_registrations_billed(self, api_client, admin_user, user_with_house, house):
        """Test that only active registrations are billed; tickets are ignored."""
        # Clean up existing data for this test
        FoodTicket.objects.filter(date__year=2025, date__month=6).delete()
        MealRegistration.objects.filter(date__year=2025, date__month=6).delete()
        self._set_zero_preferences(user_with_house)

        # Day 1: Active registration (user eats)
        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=date(2025, 6, 2),  # A Monday
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            is_active=True,
        )

        # Day 2: Active registration + ticket (registration stays active — new model)
        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=date(2025, 6, 3),  # A Tuesday
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            is_active=True,  # Registration stays active — ticket is separate
        )
        FoodTicket.objects.create(
            house=house,
            owner=user_with_house,
            date=date(2025, 6, 3),
            adults_meat=0,
            adults_veg=1,
            children_count=0,
            price=Decimal("26.00"),
            is_available=True,
        )

        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2025-06-01", "end_date": "2025-06-30"})

        assert response.status_code == 200
        data = response.json()

        owner_house_cost = next((h for h in data["houses"] if h["house_id"] == house.id), None)

        # Billing: only registrations count → 26 + 26 = 52 DKK (not 78)
        assert owner_house_cost is not None
        assert Decimal(owner_house_cost["total_cost"]) == Decimal("52.00")
        assert owner_house_cost["registration_count"] == 2
        assert owner_house_cost["adult_meat_portions"] >= 0


# =============================================================================
# Registration Lock Tests
# =============================================================================


@pytest.mark.django_db
class TestRegistrationLock:
    """Tests for registration lock enforcement after deadline."""

    def test_cannot_create_registration_after_deadline(
        self, api_client, user_with_house, monday_date
    ):
        """POST to registration-list returns 400 after the deadline."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday

        # Mock time to be AFTER the deadline (Wednesday 23:59:59 of previous week)
        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))  # Thursday
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:registration-list")
            data = {
                "date": meal_date.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
            }
            response = api_client.post(url, data, format="json")

        assert response.status_code == 400
        assert "frist" in str(response.json()).lower()

    def test_cannot_change_portions_after_deadline(self, api_client, user_with_house, monday_date):
        """PATCH adults_veg returns 400 after deadline."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday
        registration = MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))  # Thursday
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:registration-detail", kwargs={"pk": registration.pk})
            response = api_client.patch(url, {"adults_veg": 3}, format="json")

        assert response.status_code == 400

    def test_cannot_change_is_active_after_deadline(self, api_client, user_with_house, monday_date):
        """PATCH is_active returns 400 after deadline."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday
        registration = MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:registration-detail", kwargs={"pk": registration.pk})
            response = api_client.patch(url, {"is_active": False}, format="json")

        assert response.status_code == 400

    def test_can_change_dining_option_after_deadline(
        self, api_client, user_with_house, monday_date
    ):
        """PATCH dining_option succeeds after deadline."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday
        registration = MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            children_count=0,
            dining_option="eat_in",
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:registration-detail", kwargs={"pk": registration.pk})
            response = api_client.patch(url, {"dining_option": "take_away"}, format="json")

        assert response.status_code == 200
        registration.refresh_from_db()
        assert registration.dining_option == "take_away"

    def test_can_change_seating_time_after_deadline(self, api_client, user_with_house, monday_date):
        """PATCH seating_time succeeds after deadline."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday
        registration = MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            children_count=0,
            seating_time="17:30",
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:registration-detail", kwargs={"pk": registration.pk})
            response = api_client.patch(url, {"seating_time": "18:30"}, format="json")

        assert response.status_code == 200
        registration.refresh_from_db()
        assert registration.seating_time == "18:30"


@pytest.mark.django_db
class TestTicketValidationWithExisting:
    """Tests for ticket validation that accounts for existing available tickets."""

    def test_cannot_oversell_with_existing_tickets(self, api_client, user_with_house):
        """Second ticket exceeding registration minus existing tickets → 400."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday
        MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        # First ticket: sell 1 veg
        FoodTicket.objects.create(
            house=user_with_house.house,
            owner=user_with_house,
            date=meal_date,
            adults_veg=1,
            children_count=0,
            is_available=True,
        )

        # Second ticket: try to sell 2 more veg (1 available left, 2 requested)
        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:ticket-list")
            data = {
                "date": meal_date.isoformat(),
                "adults_veg": 2,
                "children_count": 0,
            }
            response = api_client.post(url, data, format="json")

        assert response.status_code == 400

    def test_multiple_partial_tickets_allowed(self, api_client, user_with_house):
        """Sum of tickets within registration bounds → 201."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday
        MealRegistration.objects.create(
            house=user_with_house.house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=3,
            children_count=0,
            is_active=True,
        )

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:ticket-list")

            # Sell 1 first
            response1 = api_client.post(
                url,
                {"date": meal_date.isoformat(), "adults_veg": 1, "children_count": 0},
                format="json",
            )
            assert response1.status_code == 201

            # Sell 1 more (total 2, within 3 limit)
            response2 = api_client.post(
                url,
                {"date": meal_date.isoformat(), "adults_veg": 1, "children_count": 0},
                format="json",
            )
            assert response2.status_code == 201

    def test_ticket_requires_active_registration(self, api_client, user_with_house):
        """Creating a ticket without an active registration returns 400."""
        api_client.force_authenticate(user=user_with_house)

        meal_date = date(2025, 12, 22)  # Monday
        # No registration created

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_now = timezone.make_aware(timezone.datetime(2025, 12, 18, 10, 0))
            mock_tz.now.return_value = mock_now
            mock_tz.get_current_timezone.return_value = timezone.get_current_timezone()

            url = reverse("food:ticket-list")
            data = {
                "date": meal_date.isoformat(),
                "adults_veg": 1,
                "children_count": 0,
            }
            response = api_client.post(url, data, format="json")

        assert response.status_code == 400


@pytest.mark.django_db
class TestStatsTicketHandling:
    """Stats totals are gross registrations — tickets never reduce them."""

    def test_stats_ignore_available_tickets(self, authenticated_client, user_with_house, house):
        """Available (unsold) tickets do NOT reduce stats totals.

        The seller is still billed for their registration whether or not the
        ticket sells, so an unsold ticket must not shrink the aggregate.
        """
        meal_date = date(2025, 12, 22)  # Monday

        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        # Create an unsold ticket for 1 portion
        FoodTicket.objects.create(
            house=house,
            owner=user_with_house,
            date=meal_date,
            adults_veg=1,
            children_count=0,
            is_available=True,  # Unsold
        )

        url = reverse("food:registration-stats")
        response = authenticated_client.get(url, {"date": meal_date.isoformat()})
        assert response.status_code == 200
        data = response.json()

        # Unsold ticket ignored → gross 2 (not 1)
        assert data["total_registrations"]["adults_veg"] == 2
        assert data["total_registrations"]["adults"] == 2
        # The old ticket-adjusted "effective" total is gone entirely.
        assert "total" not in data

    def test_stats_ignore_claimed_tickets(
        self, authenticated_client, user_with_house, house, admin_user
    ):
        """Claimed (sold) tickets are NOT subtracted — claimer is eating."""
        meal_date = date(2025, 12, 22)  # Monday

        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        # Create a claimed ticket — the claimer is eating in place of seller
        FoodTicket.objects.create(
            house=house,
            owner=user_with_house,
            date=meal_date,
            adults_veg=1,
            children_count=0,
            is_available=False,  # Claimed/sold
            claimed_by=admin_user,
            claimed_at=timezone.now(),
        )

        url = reverse("food:registration-stats")
        response = authenticated_client.get(url, {"date": meal_date.isoformat()})
        assert response.status_code == 200
        data = response.json()

        # Claimed ticket not subtracted → still 2
        assert data["total_registrations"]["adults_veg"] == 2
        assert data["total_registrations"]["adults"] == 2

    def test_stats_one_registration_per_house(
        self, authenticated_client, user_with_house, house, admin_user
    ):
        """One registration per house per date (unique constraint: house+date).

        The unique constraint now prevents duplicate registrations per house.
        """
        meal_date = date(2025, 12, 22)  # Monday

        # One registration for the house
        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=meal_date,
            adults_veg=2,
            is_active=True,
        )

        url = reverse("food:registration-stats")
        response = authenticated_client.get(url, {"date": meal_date.isoformat()})
        assert response.status_code == 200
        data = response.json()

        # One registration → 2 adults
        assert data["total_registrations"]["adults_veg"] == 2
        assert data["total_registrations"]["adults"] == 2


# =============================================================================
# Materialize-for-houses Tests
# =============================================================================


@pytest.mark.django_db
class TestMaterializeForHouses:
    """Tests for _materialize_for_houses from the tasks module."""

    def test_creates_registrations_for_houses_without_one(self, house):
        """Houses with active inhabitants but no registration get one created."""
        from apps.food.tasks import _materialize_for_houses

        user1 = User.objects.create_user(
            email="mat1@example.com", password="pass", first_name="Mat1", house=house
        )
        MealPreference.objects.create(
            house=house,
            last_modified_by=user1,
            day_of_week=0,
            adults_veg=2,
            adults_meat=1,
            children_count=0,
        )
        past_monday = date(2023, 10, 2)  # A Monday
        created = _materialize_for_houses([past_monday])
        assert created == 1
        reg = MealRegistration.objects.get(house=house, date=past_monday)
        assert reg.adults_veg == 2
        assert reg.adults_meat == 1
        assert reg.house == house

    def test_skips_houses_with_existing_registration(self, house):
        """Houses that already have a registration are skipped."""
        from apps.food.tasks import _materialize_for_houses

        user1 = User.objects.create_user(
            email="mat2@example.com", password="pass", first_name="Mat2", house=house
        )
        past_monday = date(2023, 10, 9)
        MealRegistration.objects.create(
            house=house, last_modified_by=user1, date=past_monday, adults_veg=5, is_active=True
        )
        created = _materialize_for_houses([past_monday])
        assert created == 0
        # Original registration unchanged
        assert MealRegistration.objects.get(house=house, date=past_monday).adults_veg == 5

    def test_skips_weekend_dates(self, house):
        """Weekend dates (weekday > 3) are skipped."""
        from apps.food.tasks import _materialize_for_houses

        User.objects.create_user(
            email="mat3@example.com", password="pass", first_name="Mat3", house=house
        )
        saturday = date(2023, 10, 7)
        created = _materialize_for_houses([saturday])
        assert created == 0

    def test_concurrent_calls_no_integrity_error(self, house):
        """Calling _materialize_for_houses twice for the same date doesn't raise."""
        from apps.food.tasks import _materialize_for_houses

        User.objects.create_user(
            email="mat4@example.com", password="pass", first_name="Mat4", house=house
        )
        past_monday = date(2023, 11, 6)
        created1 = _materialize_for_houses([past_monday])
        created2 = _materialize_for_houses([past_monday])
        assert created1 == 1
        assert created2 == 0  # Already exists, skipped by has_reg check
        assert MealRegistration.objects.filter(date=past_monday, house=house).count() == 1

    def test_defaults_when_no_preference(self, house):
        """Without a preference, defaults to house_count veg portions."""
        from apps.food.tasks import _materialize_for_houses

        User.objects.create_user(
            email="mat5a@example.com", password="pass", first_name="Mat5a", house=house
        )
        User.objects.create_user(
            email="mat5b@example.com", password="pass", first_name="Mat5b", house=house
        )
        past_monday = date(2023, 11, 13)
        _materialize_for_houses([past_monday])
        reg = MealRegistration.objects.get(date=past_monday, house=house)
        assert reg.adults_veg == 2  # 2 inhabitants
        assert reg.adults_meat == 0
        assert reg.dining_option == "eat_in"
        assert reg.seating_time == "17:30"

    def test_skips_houses_with_no_active_inhabitants(self, house):
        """Houses with no active inhabitants are skipped."""
        from apps.food.tasks import _materialize_for_houses

        user1 = User.objects.create_user(
            email="mat6@example.com", password="pass", first_name="Mat6", house=house
        )
        user1.is_active = False
        user1.save()
        past_monday = date(2023, 11, 20)
        created = _materialize_for_houses([past_monday])
        assert created == 0


@pytest.mark.django_db
class TestMaterializeWeekDateCalculation:
    """Tests for the date calculation in the periodic task."""

    def test_thursday_calculates_next_monday(self):
        """When run on Thursday, calculates next Monday correctly."""
        from apps.food.tasks import materialize_week_registrations

        thursday = date(2024, 3, 7)  # A Thursday
        with patch("apps.food.tasks.date") as mock_date:
            mock_date.today.return_value = thursday
            mock_date.side_effect = lambda *args, **kw: date(*args, **kw)
            with patch("apps.food.tasks._materialize_for_houses") as mock_mat:
                mock_mat.return_value = 0
                materialize_week_registrations()
                dates = mock_mat.call_args[0][0]
                assert dates[0] == date(2024, 3, 11)  # Next Monday
                assert dates[-1] == date(2024, 3, 14)  # Next Thursday
                assert len(dates) == 4


@pytest.mark.django_db
class TestBillingOnlyUsesRealRegistrations:
    """After materialization, billing only uses real MealRegistration rows."""

    def test_billing_no_registration_means_zero_cost(self, api_client, admin_user, house):
        """If materialization creates a zero-portion row, house is not billed."""
        user1 = User.objects.create_user(
            email="bilreal@example.com", password="pass", first_name="BR", house=house
        )
        # Set zero preferences so materialization creates inactive registrations
        for day in range(4):
            MealPreference.objects.create(
                house=house,
                last_modified_by=user1,
                day_of_week=day,
                adults_meat=0,
                adults_veg=0,
                children_count=0,
                dining_option="eat_in",
                seating_time="17:30",
            )
        api_client.force_authenticate(user=admin_user)
        url = reverse("food:monthly-food-cost")
        response = api_client.get(url, {"start_date": "2024-06-01", "end_date": "2024-06-30"})
        assert response.status_code == 200
        data = response.json()
        house_data = next(h for h in data["houses"] if h["house_id"] == house.id)
        assert house_data["total_cost"] == "0.00"


class TestRecipeSheetParser:
    """Unit tests for the recipe-sheet content parser (services/recipe_sheets).

    ``_build_entry_from_grid`` is pure (no Drive/network), so we feed it grids
    that mirror the real .xlsx layout: C1 dish name, F1 weekday, an ingredient
    table in C–F under an "Ingrediens" header, and Fremgangsmåde steps merged
    into column C.
    """

    def _service(self):
        from apps.food.services.recipe_sheets import RecipeSheetService

        return RecipeSheetService()

    def _pasta_grid(self):
        # Columns A..F (1-indexed). None marks an empty cell.
        return [
            ("Basis Alle", None, "Pasta", None, None, "Tirsdag"),  # 1
            ("Til (voksne)", None, "K/F", "Veg", "Vegansk", "Børn"),  # 2
            (120, 120, 0, 86, 0, 52),  # 3
            (None, None, "Mål", None, "Ingrediens", "Kommentar"),  # 4
            (9.0, 7.714, 7.2, "kg", "pasta", None),  # 5
            (None, None, None, None, "vand", None),  # 6
            (None, None, None, "spsk", "salt", "10 g salt pr L vand"),  # 7
            (None, None, None, None, None, None),  # 8 (blank)
            (None, None, "Fremgangsmåde", None, None, None),  # 9
            (None, None, "1. Kog pasta i ovnen", None, None, None),  # 10
            (None, None, "2. Hæld vandet fra", None, None, None),  # 11
            (None, None, None, None, None, None),  # 12 (merged continuation)
        ]

    def test_parses_name_weekday_ingredients_steps(self):
        entry = self._service()._build_entry_from_grid("Ti1", self._pasta_grid(), url="")
        assert entry is not None
        assert entry["code"] == "Ti1"
        assert entry["day"] == 1  # Tirsdag
        assert entry["index"] == 1
        assert entry["name"] == "Pasta"
        assert entry["weekday"] == "Tirsdag"
        assert entry["ingredients"] == [
            {"amount": "7.2", "unit": "kg", "name": "pasta", "comment": ""},
            {"amount": "", "unit": "", "name": "vand", "comment": ""},
            {"amount": "", "unit": "spsk", "name": "salt", "comment": "10 g salt pr L vand"},
        ]
        assert entry["steps"] == ["1. Kog pasta i ovnen", "2. Hæld vandet fra"]

    def test_non_recipe_title_returns_none(self):
        entry = self._service()._build_entry_from_grid("Indkøbsliste", self._pasta_grid(), url="")
        assert entry is None

    def test_weekday_derived_from_code_not_f1(self):
        # F1 says "Onsdag" (a stale copy-paste) but the To3 code means Torsdag.
        grid = list(self._pasta_grid())
        grid[0] = ("Basis Alle", None, "Koldhævet brød", None, None, "Onsdag")
        entry = self._service()._build_entry_from_grid("To3", grid, url="")
        assert entry is not None
        assert entry["day"] == 3
        assert entry["weekday"] == "Torsdag"

    def test_amount_column_falls_back_to_b_when_c_empty(self):
        # The "Ris" case: quantities only in column B, C is empty.
        grid = [
            ("Basis Alle", None, "Ris", None, None, "Mandag"),
            ("Til (voksne)", None, "K/F", "Veg", "Vegansk", "Børn"),
            (120, 120, 0, 86, 0, 52),
            (None, None, "Mål", None, "Ingrediens", "Kommentar"),
            (9.6, 9.6, None, "kg", "parboiled ris", None),
            (14.5, 14.5, None, "L", "vand", None),
            (None, None, "Fremgangsmåde", None, None, None),
            (None, None, "1. Skyl risene", None, None, None),
        ]
        entry = self._service()._build_entry_from_grid("Ma2", grid, url="")
        assert entry is not None
        assert [i["amount"] for i in entry["ingredients"]] == ["9.6", "14.5"]

    def test_carries_current_schema_version(self):
        from apps.food.services.recipe_sheets import CACHE_SCHEMA_VERSION

        entry = self._service()._build_entry_from_grid("Ma1", self._pasta_grid(), url="")
        assert entry is not None
        assert entry["_v"] == CACHE_SCHEMA_VERSION


class TestWeekRecipesView:
    """Tests for the standalone "Ugens opskrifter" endpoint."""

    def test_returns_cached_week_recipes_grouped_payload(self, authenticated_client, db):
        from apps.food.services.recipe_sheets import CACHE_SCHEMA_VERSION

        today = timezone.now().date()
        week_number = today.isocalendar()[1]
        year = today.isocalendar()[0]

        # Pre-populate the cache so the view serves it without any Drive call.
        DriveMenuCache.objects.create(
            week_number=week_number,
            year=year,
            drive_folder_id="folder123",
            recipe_file_id="file456",
            recipe_sheets=[
                {
                    "code": "Ti1",
                    "day": 1,
                    "index": 1,
                    "name": "Pasta",
                    "weekday": "Tirsdag",
                    "url": "",
                    "ingredients": [{"amount": "9", "unit": "kg", "name": "pasta", "comment": ""}],
                    "steps": ["1. Kog pasta"],
                    "_v": CACHE_SCHEMA_VERSION,
                },
                {
                    "code": "Ma1",
                    "day": 0,
                    "index": 1,
                    "name": "Nudler",
                    "weekday": "Mandag",
                    "url": "",
                    "ingredients": [],
                    "steps": ["1. Kog nudler"],
                    "_v": CACHE_SCHEMA_VERSION,
                },
            ],
        )

        url = reverse("food:recipes-week")
        response = authenticated_client.get(url)
        assert response.status_code == 200
        data = response.json()
        assert data["week_number"] == week_number
        assert data["year"] == year
        assert "drive.google.com/drive/folders/folder123" in data["recipe_folder_url"]
        assert "file456" in data["recipe_file_url"]
        # Sorted by (day, index): Monday before Tuesday.
        assert [r["name"] for r in data["recipes"]] == ["Nudler", "Pasta"]
        assert data["recipes"][1]["ingredients"][0]["name"] == "pasta"

    def test_no_cache_returns_empty_recipes(self, authenticated_client, db):
        url = reverse("food:recipes-week")
        response = authenticated_client.get(url, {"week": 2, "year": 2099})
        assert response.status_code == 200
        data = response.json()
        assert data["week_number"] == 2
        assert data["year"] == 2099
        assert data["recipes"] == []

    def test_requires_authentication(self, api_client, db):
        url = reverse("food:recipes-week")
        response = api_client.get(url)
        assert response.status_code in (401, 403)


class TestFrontPageParser:
    """Unit tests for the per-day 'Dagens forside' .docx parser."""

    def _build_doc(self):
        from docx import Document
        from docx.enum.text import WD_BREAK

        doc = Document()
        # Overview page: lists multiple weekdays -> must be skipped.
        doc.add_paragraph("Uge test")
        doc.add_paragraph("Mandag")
        doc.add_paragraph("Tirsdag")
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        # Mandag detail page.
        doc.add_paragraph("Mandag")
        doc.add_paragraph("Ret M")  # first content line -> title
        doc.add_paragraph("Tilbehør")  # known section label -> heading
        doc.add_paragraph("Ris og brød")
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        # Tirsdag detail page.
        doc.add_paragraph("Tirsdag")
        doc.add_paragraph("Ret T")
        doc.add_paragraph("Ved servering")  # known section label -> heading
        doc.add_paragraph("På fad")
        return doc

    def test_extracts_per_day_sections(self):
        from apps.food.services.drive_menu import FRONT_PAGE_SCHEMA_VERSION, DriveMenuService

        pages = DriveMenuService()._parse_front_pages(self._build_doc())
        assert [p["day"] for p in pages] == [0, 1]

        mon = pages[0]
        assert mon["weekday"] == "Mandag"
        assert mon["title"] == "Ret M"
        assert mon["blocks"] == [
            {"text": "Tilbehør", "heading": True},
            {"text": "Ris og brød", "heading": False},
        ]
        assert mon["_v"] == FRONT_PAGE_SCHEMA_VERSION

        tue = pages[1]
        assert tue["title"] == "Ret T"
        assert tue["blocks"][0] == {"text": "Ved servering", "heading": True}

    def test_overview_only_doc_yields_nothing(self):
        from docx import Document

        from apps.food.services.drive_menu import DriveMenuService

        doc = Document()
        doc.add_paragraph("Uge test")
        doc.add_paragraph("Mandag")
        doc.add_paragraph("Tirsdag")
        doc.add_paragraph("Onsdag")
        doc.add_paragraph("Torsdag")
        assert DriveMenuService()._parse_front_pages(doc) == []


class TestWeekRecipesFrontPages:
    """The week endpoint surfaces cached per-day front pages."""

    def test_week_endpoint_returns_cached_front_pages(self, authenticated_client, db):
        from apps.food.services.drive_menu import FRONT_PAGE_SCHEMA_VERSION

        DriveMenuCache.objects.create(
            week_number=11,
            year=2099,
            daily_front_pages=[
                {
                    "day": 1,
                    "weekday": "Tirsdag",
                    "title": "Green curry",
                    "blocks": [{"text": "Tilbehør", "heading": True}],
                    "_v": FRONT_PAGE_SCHEMA_VERSION,
                }
            ],
        )
        url = reverse("food:recipes-week")
        response = authenticated_client.get(url, {"week": 11, "year": 2099})
        assert response.status_code == 200
        data = response.json()
        assert len(data["front_pages"]) == 1
        assert data["front_pages"][0]["title"] == "Green curry"
        assert data["front_pages"][0]["blocks"][0]["heading"] is True


# =============================================================================
# import_madtilmeldinger: opt-out handling
# =============================================================================


@pytest.mark.django_db
class TestImportMadtilmeldingerOptOut:
    """The importer must record opt-out (zero) days as inactive tombstones.

    Post-deadline, _materialize_for_houses recreates any *missing* registration
    from the house's standing MealPreference. If the importer skipped opt-out
    days (leaving no row), that safety net would resurrect a meal the house
    deliberately cancelled. Writing an inactive row blocks that.
    """

    def _make_house_with_inhabitant(self):
        from apps.houses.models import House

        house = House.objects.create(name="Kløverbakkevej 49", address="x")
        assert house.slug == "49"
        User.objects.create_user(email="h49@example.com", password="x", first_name="A", house=house)
        return house

    def _row_mon_active_wed_optout(self):
        # husnr, Mon[V,W,ch], Tue[V,W,ch], Wed[V,W,KF,ch], Thu[V,W,ch]
        return ["49", "2", "0", "0", "2", "0", "0", "0", "0", "0", "0", "2", "0", "0"]

    def test_optout_day_written_as_inactive_tombstone(self):
        from apps.food.management.commands.import_madtilmeldinger import Command

        house = self._make_house_with_inhabitant()
        monday = date(2026, 6, 1)
        wednesday = date(2026, 6, 3)

        # Simulate a pre-existing (spurious) active Wednesday registration.
        MealRegistration.objects.create(
            house=house, date=wednesday, adults_veg=0, adults_meat=2, is_active=True
        )

        parsed = {"house_data": [self._row_mon_active_wed_optout()]}
        Command()._import_registrations(parsed, {"49": house}, monday, [])

        # Active day imported normally.
        mon = MealRegistration.objects.get(house=house, date=monday)
        assert mon.is_active is True
        assert (mon.adults_veg, mon.adults_meat) == (2, 0)

        # Opt-out Wednesday is NOT deleted — it becomes an inactive tombstone.
        wed = MealRegistration.objects.get(house=house, date=wednesday)
        assert wed.is_active is False
        assert (wed.adults_veg, wed.adults_meat, wed.children_count) == (0, 0, 0)

    def test_tombstone_blocks_rematerialization(self):
        from apps.food.management.commands.import_madtilmeldinger import Command
        from apps.food.tasks import _materialize_for_houses

        house = self._make_house_with_inhabitant()
        monday = date(2026, 6, 1)
        wednesday = date(2026, 6, 3)

        # House's standing order *includes* Wednesday meat — the resurrection trap.
        MealPreference.objects.create(
            house=house, day_of_week=DayOfWeek.WEDNESDAY, adults_veg=0, adults_meat=2
        )

        parsed = {"house_data": [self._row_mon_active_wed_optout()]}
        Command()._import_registrations(parsed, {"49": house}, monday, [])

        # The post-deadline materialization net runs over the opt-out date.
        _materialize_for_houses([wednesday])

        # It must NOT resurrect the cancelled Wednesday meal.
        wed = MealRegistration.objects.get(house=house, date=wednesday)
        assert wed.is_active is False
        assert (wed.adults_veg, wed.adults_meat) == (0, 0)
        assert not MealRegistration.objects.filter(date=wednesday, is_active=True).exists()

    def _write_week_csv(self, tmp_path, week: int):
        """Write a minimal single-week sheet CSV in the madtilmeldinger format."""
        lines = [
            "V=Vegetar W=Weganer KF=Kød/Fisk Børn (1-11)",
            ",,,,,,,,,,,,,,,",  # row 1: holiday notes (none)
            "Uge,Mandag (Vegetar),,,Tirsdag (Vegetar),,,Onsdag (Kød),,,,Torsdag (Vegetar),,,,",
            f"{week},Voksne,,Børn,Voksne,,Børn,Voksne,,,Børn,Voksne,,Børn,,",
            "Husnr,V,W,1-11,V,W,1-11,V,W,KF,1-11,V,W,1-11,Husnr,Pris",
            "49,2,0,0,2,0,0,0,0,0,0,2,0,0,49,156",
        ]
        path = tmp_path / f"uge{week}.csv"
        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    def test_exclude_weeks_skips_that_week(self, tmp_path):
        from django.core.management import call_command

        self._make_house_with_inhabitant()
        source = self._write_week_csv(tmp_path, week=23)

        call_command("import_madtilmeldinger", source, year=2026, exclude_weeks=[23])

        # Excluded → nothing written for week 23 (Mon 2026-06-01).
        assert not MealRegistration.objects.filter(date=date(2026, 6, 1)).exists()

    def test_without_exclude_week_is_imported(self, tmp_path):
        from django.core.management import call_command

        self._make_house_with_inhabitant()
        source = self._write_week_csv(tmp_path, week=23)

        call_command("import_madtilmeldinger", source, year=2026)

        # Monday active, Wednesday opt-out tombstoned.
        assert MealRegistration.objects.get(date=date(2026, 6, 1)).is_active is True
        assert MealRegistration.objects.get(date=date(2026, 6, 3)).is_active is False


# =============================================================================
# Meal Price Tests
# =============================================================================


@pytest.mark.django_db
class TestMealPriceSchedule:
    """Prices are resolved by meal date, so past billing never changes."""

    def test_seeded_schedule_covers_history_and_increase(self):
        from apps.food.pricing import get_prices

        # Anything before 2026-08-02 keeps the original prices.
        old = get_prices(date(2026, 8, 1))
        assert (old.adult_meat, old.adult_veg, old.child) == (
            Decimal("37.00"),
            Decimal("26.00"),
            Decimal("18.00"),
        )

        # From the cutover date (inclusive) the new prices apply.
        new = get_prices(date(2026, 8, 2))
        assert (new.adult_meat, new.adult_veg, new.child) == (
            Decimal("40.00"),
            Decimal("30.00"),
            Decimal("18.00"),
        )

    def test_dates_before_the_baseline_fall_back(self):
        from apps.food.pricing import FALLBACK_PRICES, get_prices

        assert get_prices(date(1999, 1, 1)) == FALLBACK_PRICES

    def test_total_uses_all_three_prices(self):
        from apps.food.pricing import get_prices

        # 1 * 40 + 2 * 30 + 3 * 18 = 154
        assert get_prices(date(2026, 8, 5)).total(1, 2, 3) == Decimal("154.00")

    def test_schedule_resolves_latest_matching_set(self, admin_user):
        from apps.food.models import MealPrice
        from apps.food.pricing import get_price_schedule

        MealPrice.objects.create(
            effective_from=date(2027, 1, 1),
            price_adult_meat=Decimal("50.00"),
            price_adult_veg=Decimal("40.00"),
            price_child=Decimal("25.00"),
            created_by=admin_user,
        )
        schedule = get_price_schedule()

        assert schedule.for_date(date(2026, 12, 31)).adult_meat == Decimal("40.00")
        assert schedule.for_date(date(2027, 1, 1)).adult_meat == Decimal("50.00")
        assert schedule.for_date(date(2030, 6, 1)).adult_meat == Decimal("50.00")


@pytest.mark.django_db
class TestMealPriceBilling:
    """Cost reports must price each meal at the rates of the day it was served."""

    @staticmethod
    def _zero_preferences(user):
        for day in range(4):
            MealPreference.objects.get_or_create(
                house=user.house,
                day_of_week=day,
                defaults={
                    "adults_meat": 0,
                    "adults_veg": 0,
                    "children_count": 0,
                    "dining_option": "eat_in",
                    "seating_time": "17:30",
                    "last_modified_by": user,
                },
            )

    def test_report_spanning_the_cutover_mixes_prices(
        self, api_client, admin_user, user_with_house, house
    ):
        self._zero_preferences(user_with_house)

        # Mon 2026-07-27 (old prices) and Mon 2026-08-03 (new prices).
        for reg_date in (date(2026, 7, 27), date(2026, 8, 3)):
            MealRegistration.objects.create(
                house=house,
                last_modified_by=user_with_house,
                date=reg_date,
                adults_meat=0,
                adults_veg=2,
                children_count=1,
                is_active=True,
            )

        api_client.force_authenticate(user=admin_user)
        response = api_client.get(
            reverse("food:monthly-food-cost"),
            {"start_date": "2026-07-27", "end_date": "2026-08-06"},
        )

        assert response.status_code == 200
        house_cost = next(h for h in response.json()["houses"] if h["house_id"] == house.id)
        # Old week: 2 * 26 + 18 = 70. New week: 2 * 30 + 18 = 78. Total 148.
        assert Decimal(house_cost["total_cost"]) == Decimal("148.00")

    def test_my_expenses_uses_meal_date_prices(self, api_client, user_with_house, house):
        self._zero_preferences(user_with_house)

        MealRegistration.objects.create(
            house=house,
            last_modified_by=user_with_house,
            date=date(2026, 7, 27),
            adults_meat=0,
            adults_veg=2,
            children_count=0,
            is_active=True,
        )

        api_client.force_authenticate(user=user_with_house)
        response = api_client.get(
            reverse("food:my-expenses"),
            {"start_date": "2026-07-27", "end_date": "2026-07-30"},
        )

        assert response.status_code == 200
        # 2 veg @ the old 26 DKK — unaffected by the August increase.
        assert Decimal(response.json()["total_cost"]) == Decimal("52.00")


@pytest.mark.django_db
class TestMealPriceAPI:
    """The price schedule is readable by all, writable only by food admins."""

    def test_list_requires_auth_only(self, api_client, user_with_house):
        api_client.force_authenticate(user=user_with_house)
        response = api_client.get(reverse("food:price-list"))

        assert response.status_code == 200
        body = response.json()
        assert len(body) == 2
        # Newest first, and prices are numbers so the UI can do arithmetic.
        assert body[0]["effective_from"] == "2026-08-02"
        assert body[0]["price_adult_meat"] == 40.0

    def test_non_admin_cannot_create(self, api_client, user_with_house):
        api_client.force_authenticate(user=user_with_house)
        response = api_client.post(
            reverse("food:price-list"),
            {
                "effective_from": (timezone.localdate() + timedelta(days=30)).isoformat(),
                "price_adult_meat": "45.00",
                "price_adult_veg": "35.00",
                "price_child": "20.00",
            },
            format="json",
        )

        assert response.status_code == 403

    def test_food_admin_can_create_future_price_set(self, api_client, admin_user):
        api_client.force_authenticate(user=admin_user)
        effective = timezone.localdate() + timedelta(days=30)
        response = api_client.post(
            reverse("food:price-list"),
            {
                "effective_from": effective.isoformat(),
                "price_adult_meat": "45.00",
                "price_adult_veg": "35.00",
                "price_child": "20.00",
                "note": "Nye råvarepriser",
            },
            format="json",
        )

        assert response.status_code == 201, response.json()
        assert response.json()["created_by_name"] == admin_user.get_full_name()

        from apps.food.pricing import get_prices

        assert get_prices(effective).adult_meat == Decimal("45.00")

    def test_cannot_backdate_a_price_set(self, api_client, admin_user):
        api_client.force_authenticate(user=admin_user)
        response = api_client.post(
            reverse("food:price-list"),
            {
                "effective_from": (timezone.localdate() - timedelta(days=1)).isoformat(),
                "price_adult_meat": "45.00",
                "price_adult_veg": "35.00",
                "price_child": "20.00",
            },
            format="json",
        )

        assert response.status_code == 400
        assert "fortiden" in str(response.json())

    def test_duplicate_start_date_is_rejected(self, api_client, admin_user):
        api_client.force_authenticate(user=admin_user)
        effective = (timezone.localdate() + timedelta(days=10)).isoformat()
        payload = {
            "effective_from": effective,
            "price_adult_meat": "45.00",
            "price_adult_veg": "35.00",
            "price_child": "20.00",
        }

        assert (
            api_client.post(reverse("food:price-list"), payload, format="json").status_code == 201
        )

        response = api_client.post(reverse("food:price-list"), payload, format="json")
        assert response.status_code == 400
        assert "allerede et prissæt" in str(response.json())

    def test_negative_price_is_rejected(self, api_client, admin_user):
        api_client.force_authenticate(user=admin_user)
        response = api_client.post(
            reverse("food:price-list"),
            {
                "effective_from": (timezone.localdate() + timedelta(days=10)).isoformat(),
                "price_adult_meat": "-5.00",
                "price_adult_veg": "35.00",
                "price_child": "20.00",
            },
            format="json",
        )

        assert response.status_code == 400

    def test_future_price_set_can_be_edited_and_deleted(self, api_client, admin_user):
        from apps.food.models import MealPrice

        api_client.force_authenticate(user=admin_user)
        obj = MealPrice.objects.create(
            effective_from=timezone.localdate() + timedelta(days=20),
            price_adult_meat=Decimal("45.00"),
            price_adult_veg=Decimal("35.00"),
            price_child=Decimal("20.00"),
            created_by=admin_user,
        )
        url = reverse("food:price-detail", args=[obj.pk])

        response = api_client.patch(url, {"price_adult_meat": "46.00"}, format="json")
        assert response.status_code == 200
        assert response.json()["price_adult_meat"] == 46.0

        assert api_client.delete(url).status_code == 204
        assert not MealPrice.objects.filter(pk=obj.pk).exists()

    def test_price_set_in_effect_cannot_be_edited_or_deleted(self, api_client, admin_user):
        from apps.food.models import MealPrice

        api_client.force_authenticate(user=admin_user)
        obj = MealPrice.objects.get(effective_from=date(2026, 8, 2))
        url = reverse("food:price-detail", args=[obj.pk])

        with patch("apps.food.serializers.timezone") as mock_tz:
            mock_tz.localdate.return_value = date(2026, 9, 1)
            response = api_client.patch(url, {"price_adult_meat": "99.00"}, format="json")
        assert response.status_code == 400
        assert "trådt i kraft" in str(response.json())

        with patch("apps.food.views.timezone") as mock_tz:
            mock_tz.localdate.return_value = date(2026, 9, 1)
            response = api_client.delete(url)
        assert response.status_code == 400
        assert MealPrice.objects.filter(pk=obj.pk).exists()

    def test_last_price_set_cannot_be_deleted(self, api_client, admin_user):
        from apps.food.models import MealPrice

        api_client.force_authenticate(user=admin_user)
        MealPrice.objects.exclude(effective_from=date(2026, 8, 2)).delete()
        obj = MealPrice.objects.get()

        response = api_client.delete(reverse("food:price-detail", args=[obj.pk]))
        assert response.status_code == 400
        assert MealPrice.objects.filter(pk=obj.pk).exists()
